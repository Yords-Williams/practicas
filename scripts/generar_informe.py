"""Genera INFORME_FINAL_PRACTICAS.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
AZUL_OSCURO = RGBColor(0x1F, 0x2A, 0x5E)
AZUL_MEDIO  = RGBColor(0x2E, 0x74, 0xB5)

def centrado(text, bold=False, size=12, color=None, espacio_antes=0, espacio_despues=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after  = Pt(espacio_despues)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def titulo(text, level=1, color=AZUL_OSCURO):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return h

def parrafo(text, size=11, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def separador():
    p = doc.add_paragraph('─' * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

def tabla_simple(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # encabezado
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # fondo azul
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
    # filas
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════
centrado('UNIVERSIDAD NACIONAL DE JULIACA', bold=True, size=16, color=AZUL_OSCURO, espacio_antes=20)
centrado('Facultad de Ciencias de Ingenierías', bold=True, size=13, color=AZUL_OSCURO)
centrado('Escuela Profesional de Ingeniería de Software y Sistemas', size=12, color=AZUL_OSCURO, espacio_despues=30)

doc.add_paragraph()
centrado('INFORME FINAL DE PRÁCTICAS PREPROFESIONALES', bold=True, size=15, color=AZUL_OSCURO, espacio_antes=20, espacio_despues=20)
doc.add_paragraph()

datos_portada = [
    ('Institución:',           'Municipalidad Distrital de Caracoto'),
    ('Campo ocupacional:',     'Practicante'),
    ('Período cubierto:',      'Del 16 de marzo al 22 de mayo del 2026 (Semanas 1–10)'),
    ('Practicante:',           'Yords Williams Ccalla Mamani'),
    ('DNI:',                   '75093371'),
    ('Código de estudiante:',  '2022107039'),
    ('Semestre concluido:',    'Octavo'),
    ('Fecha de elaboración:',  '15 de julio del 2026'),
]
t = doc.add_table(rows=len(datos_portada), cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(datos_portada):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[1].text = v
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[0].width = Cm(5)
    t.rows[i].cells[1].width = Cm(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
titulo('1. INTRODUCCIÓN')
parrafo(
    'El presente informe documenta de manera detallada las actividades realizadas durante las primeras diez semanas '
    'del período de prácticas preprofesionales, llevadas a cabo en la Municipalidad Distrital de Caracoto, Provincia '
    'de San Román, Departamento de Puno, en el marco de la formación académica de la Escuela Profesional de '
    'Ingeniería de Software y Sistemas de la Universidad Nacional de Juliaca.'
)
parrafo(
    'El proyecto desarrollado consiste en un Sistema de Monitoreo Inteligente denominado "CCTV AI PRO", cuyo '
    'propósito es apoyar a la Unidad de Seguridad Ciudadana y el Serenazgo Municipal de Caracoto mediante la '
    'detección automática de incidencias críticas en tiempo real: incendios, robos menores y choques de vehículos. '
    'El sistema integra cámaras IP/RTSP existentes en la institución con modelos de inteligencia artificial basados '
    'en visión por computadora, específicamente el detector YOLOv8 de Ultralytics.'
)
parrafo(
    'La motivación principal de esta práctica reside en la necesidad real de la municipalidad de modernizar su '
    'capacidad de vigilancia pública, pasando de una supervisión manual —sujeta a fatiga y error humano— a un '
    'esquema automatizado que genera alertas inmediatas, permitiendo una respuesta más rápida y eficiente del '
    'personal operativo ante situaciones de riesgo. Esta implementación aplica directamente los conocimientos '
    'adquiridos en las áreas de programación, inteligencia artificial, diseño de software y bases de datos '
    'cursadas durante la carrera.'
)
parrafo(
    'El presente documento cubre las Actividades 1 a 8 del Plan de Trabajo (Semanas 1 a 10), desde el análisis '
    'de requerimientos hasta la finalización del módulo de detección de choques de autos, conforme al cronograma '
    'establecido. El informe se estructura en: Objetivos, Resumen de actividades semanales, Análisis FODA, '
    'Sugerencias y Conclusiones.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. OBJETIVOS
# ══════════════════════════════════════════════════════════════════════════════
titulo('2. OBJETIVOS')

titulo('2.1. Objetivo General', level=2)
parrafo(
    'Desarrollar y aplicar los conocimientos y experiencias adquiridos en el transcurso de la formación en la '
    'Escuela Profesional de Ingeniería de Software y Sistemas, optimizando procesos, realizando análisis, diseño '
    'e implementación de soluciones tecnológicas dentro del centro de prácticas, específicamente mediante la '
    'construcción de un sistema de monitoreo inteligente basado en visión por computadora para la Municipalidad '
    'Distrital de Caracoto.'
)

titulo('2.2. Objetivos Específicos', level=2)
bullet(
    'Implementar un sistema de monitoreo inteligente para la detección de ocurrencias (incendios, robos menores '
    'y choques de vehículos) para la Unidad de Seguridad Ciudadana y el Serenazgo Municipal de Caracoto, '
    'integrando modelos de detección YOLOv8 sobre flujos de video en tiempo real.'
)
bullet(
    'Diseñar e implementar una interfaz gráfica funcional (GUI) con PySide6 que permita al personal operativo '
    'visualizar múltiples cámaras simultáneamente, cambiar modos de detección y gestionar incidentes registrados.'
)
bullet(
    'Integrar los módulos de detección (incendios, robos y choques) con un sistema centralizado de base de datos '
    '(SQLite) y un módulo de alarmas, garantizando la interoperabilidad y cohesión del sistema completo.'
)
bullet(
    'Brindar soporte técnico eficiente a los empleados de la institución, identificando y solucionando problemas '
    'relacionados con hardware, software y redes dentro del área de trabajo asignada.'
)
bullet(
    'Cumplir con las actividades del Plan de Trabajo y demás tareas encomendadas por el jefe inmediato del '
    'centro de prácticas dentro de los plazos establecidos.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. RESUMEN POR SEMANA
# ══════════════════════════════════════════════════════════════════════════════
titulo('3. RESUMEN DE LAS ACCIONES Y/O ACTIVIDADES REALIZADAS POR SEMANA')
parrafo(
    'El siguiente apartado presenta de forma cronológica y cualitativa las actividades desempeñadas durante las '
    'primeras diez semanas del período de prácticas, cubriendo desde el análisis de requerimientos hasta la '
    'culminación del módulo de detección de choques de autos.'
)

semanas = [
    {
        'titulo': 'SEMANA 1 — 16 al 20 de marzo de 2026',
        'actividades_plan': 'Análisis de requerimientos · Investigación',
        'horas': '15 horas (3 h/día)',
        'actividad_principal': 'Inicio del análisis de requerimientos e investigación bibliográfica.',
        'acciones': [
            'Reunión de presentación con el jefe de la Unidad de Seguridad Ciudadana y el responsable de tecnología de la Municipalidad Distrital de Caracoto.',
            'Realización de entrevistas con el personal operativo del serenazgo para identificar las principales problemáticas: respuesta tardía ante incendios, dificultad para detectar robos en zonas ciegas y falta de registro automatizado de accidentes viales.',
            'Levantamiento de requerimientos funcionales: detección de incendios, robos menores y choques; visualización en tiempo real de múltiples cámaras; sistema de alertas automáticas.',
            'Levantamiento de requerimientos no funcionales: tiempo de respuesta < 2 s, compatibilidad con cámaras RTSP existentes, operación en hardware disponible (NVIDIA Quadro P1000, Intel Core i7-10700).',
            'Inicio de revisión bibliográfica sobre sistemas de visión por computadora aplicados a seguridad ciudadana.',
            'Identificación de modelos candidatos: YOLOv8 (Ultralytics), OpenPose, DeepSORT.',
        ],
        'resultados': 'Documento de requerimientos preliminar con 12 requerimientos funcionales y 8 no funcionales. Inventario de hardware y 4 cámaras IP con protocolo RTSP.',
    },
    {
        'titulo': 'SEMANA 2 — 23 al 27 de marzo de 2026',
        'actividades_plan': 'Análisis de requerimientos · Investigación · Diseño del sistema (inicio)',
        'horas': '15 horas',
        'actividad_principal': 'Cierre de requerimientos, investigación técnica e inicio del diseño de arquitectura.',
        'acciones': [
            'Consolidación y validación del documento de requerimientos con el jefe inmediato; aprobación de los requerimientos funcionales definitivos.',
            'Evaluación comparativa de modelos YOLO: YOLOv8n, YOLOv8s y YOLOv8m. Se seleccionó YOLOv8n por su equilibrio entre velocidad y precisión en hardware con 2 GB de VRAM.',
            'Revisión de estudios de caso sobre detección de incendios con YOLO (dataset UCF-Arson, Fire-8) y datasets de accidentes de tráfico.',
            'Evaluación de frameworks de GUI: PyQt5, PySide6 y Tkinter. Se seleccionó PySide6 por su licencia LGPL, soporte en Windows y rendimiento en renderizado de video.',
            'Inicio del diseño de la arquitectura general del sistema: identificación de módulos principales.',
            'Definición del stack tecnológico: Python 3.11, PyTorch 2.x con CUDA 11.8, Ultralytics YOLOv8, OpenCV 4.x, PySide6, SQLite.',
        ],
        'resultados': 'Documento de requerimientos v1.0 aprobado. Stack tecnológico seleccionado y documentado con justificación técnica.',
    },
    {
        'titulo': 'SEMANA 3 — 30 de marzo al 3 de abril de 2026',
        'actividades_plan': 'Investigación · Diseño del sistema · Diseño GUI (inicio)',
        'horas': '15 horas',
        'actividad_principal': 'Diseño de arquitectura modular y primeros wireframes de la interfaz gráfica.',
        'acciones': [
            'Finalización de la investigación bibliográfica con revisión de Goodfellow et al. (2016), Sommerville (2016) y Pressman & Maxim (2014).',
            'Diseño de la arquitectura modular: camera.py/camera_dialog.py (captura RTSP), detector.py (motor YOLO general), modulos/fire_detector.py, modulos/accident_detection.py, modulos/robo_detector.py, modulos/alarma/ (notificaciones), ui.py (GUI), config.py (configuración centralizada).',
            'Diseño del diagrama de flujo del pipeline: captura de frame → redimensionado → detección según modo → registro de incidente → notificación.',
            'Diseño del esquema de base de datos SQLite: tabla incidents (id, camera_name, incident_type, severity, confidence, details, created_at, status, image_path) y tabla recipients.',
            'Inicio de wireframes de la GUI: grilla de cámaras 2×3, barra lateral de navegación, panel de estadísticas y tabla de incidentes.',
        ],
        'resultados': 'Diagrama de arquitectura del sistema aprobado. Wireframes iniciales de la GUI documentados.',
    },
    {
        'titulo': 'SEMANA 4 — 7 al 11 de abril de 2026',
        'actividades_plan': 'Diseño del sistema · Diseño GUI',
        'horas': '15 horas',
        'actividad_principal': 'Finalización del diseño del sistema y diseño detallado de la interfaz gráfica.',
        'acciones': [
            'Definición completa de la estructura de base de datos SQLite y archivo de configuración fire_config.json.',
            'Diseño detallado de la GUI: vista de cámaras (grilla 3×2, hasta 6 cámaras), selector de modo (General / Accidentes / Personas / Robos), vista de incidentes con exportación a Excel, gestión de destinatarios y panel de estadísticas.',
            'Diseño de la barra lateral estilo dashboard: ítems Dashboard, Cámaras, Incidentes, Destinatarios, Notificaciones y Ajustes.',
            'Definición del archivo cameras_config.json para persistencia de la configuración de cámaras con ajustes por cámara (detect_fire, detect_theft, detect_accident).',
            'Revisión y aprobación del diseño completo con el jefe inmediato.',
        ],
        'resultados': 'Diseño funcional completo del sistema aprobado. Especificación de GUI lista para implementación.',
    },
    {
        'titulo': 'SEMANA 5 — 14 al 18 de abril de 2026',
        'actividades_plan': 'Diseño GUI · Configuración del entorno · Captura de video · Detección de incendios (inicio)',
        'horas': '15 horas',
        'actividad_principal': 'Configuración del entorno de desarrollo e inicio de la implementación.',
        'acciones': [
            'Creación del entorno virtual Python 3.11 e instalación de todas las dependencias: ultralytics, torch+CUDA 11.8, opencv-python, PySide6, numpy, pillow, openpyxl.',
            'Configuración de variables de entorno para optimización: OPENBLAS_NUM_THREADS=1, OMP_NUM_THREADS=2, MKL_NUM_THREADS=2.',
            'Verificación de GPU NVIDIA Quadro P1000 (2 GB VRAM) con CUDA 11.8. Descarga del modelo base yolov8n.pt.',
            'Implementación del módulo de captura de video (camera.py): clase CamaraRTSP con hilo independiente, reconexión automática y método read() thread-safe.',
            'Implementación del diálogo de gestión de cámaras (camera_dialog.py) con persistencia en cameras_config.json.',
            'Estructura inicial de FireDetectionSystem con inicialización automática de la base de datos SQLite.',
        ],
        'resultados': 'Entorno de desarrollo completamente funcional. Módulo de captura de video operativo y probado con cámara local.',
    },
    {
        'titulo': 'SEMANA 6 — 21 al 25 de abril de 2026',
        'actividades_plan': 'Captura de video · Detección de incendios · Detección de robos (inicio)',
        'horas': '15 horas',
        'actividad_principal': 'Implementación completa del módulo de incendios y primeras pruebas con video real.',
        'acciones': [
            'Pruebas de conexión con las cámaras RTSP de la municipalidad. Verificación de latencia < 150 ms en red local.',
            'Implementación de detect_fire(): detección híbrida con modelo YOLO (incendio/best.pt) como detector principal y análisis de color HSV como respaldo.',
            'Implementación de analyze_frame(): gestión del cooldown entre alertas (60 s), registro del incidente en SQLite y notificación a destinatarios.',
            'Implementación de add_recipient(), list_recipients(), save_incident() y send_report() para gestión completa del ciclo de alerta.',
            'Configuración externa en fire_config.json: threshold (0.12), min_area (180 px), cooldown (60 s).',
            'Estructura inicial del módulo robo_detector.py: definición de KalmanBoxTracker y SortTracker para tracking de personas.',
        ],
        'resultados': 'FireDetectionSystem funcional con detección híbrida YOLO+HSV. Tasa de detección estimada: 87% en escenas con llamas visibles.',
    },
    {
        'titulo': 'SEMANA 7 — 28 de abril al 2 de mayo de 2026',
        'actividades_plan': 'Detección de incendios · Detección de robos · Detección de choques (inicio)',
        'horas': '15 horas',
        'actividad_principal': 'Integración del módulo de incendios en la GUI, avance en robos e inicio de accident_detection.',
        'acciones': [
            'Integración de fire_detector.py en ui.py: el detector de incendios se ejecuta en cada frame de todas las cámaras habilitadas, independientemente del modo seleccionado.',
            'Implementación de detect_people_and_objects(): detección de personas y objetos de valor (mochilas, maletines, laptops, celulares, cuchillos) con umbrales diferenciados.',
            'Implementación de track_people() con SortTracker (filtro de Kalman + asignación húngara por IoU) y PersonAppearanceTracker (re-identificación por histograma HSV).',
            'Implementación de match_objects_to_people(): asignación de objetos a personas por proximidad euclídea.',
            'Inicio de accident_detection.py: constructor con rutas de modelos configurables (choques/best.pt y detector_de_auto_con_dano.pt).',
            'Implementación inicial de detect_vehicles(): filtrado de clases COCO car(2), motorcycle(3), bus(5), truck(7).',
        ],
        'resultados': 'Módulo de incendios integrado en la UI. Pipeline básico de detección de personas y objetos funcional. Estructura del módulo de choques iniciada.',
    },
    {
        'titulo': 'SEMANA 8 — 5 al 9 de mayo de 2026',
        'actividades_plan': 'Detección de robos menores · Detección de choques de autos',
        'horas': '15 horas',
        'actividad_principal': 'Avance significativo en el pipeline de robos e implementación del núcleo de accident_detection.',
        'acciones': [
            'Implementación de detect_suspicious_activity(): análisis de movimientos rápidos y posturas para inferir actividad sospechosa.',
            'Implementación de analyze_theft(): scoring ponderado con transferencia de objetos (40%), presencia de armas (30%) y actividad sospechosa (30%); umbral de alarma: 0.70.',
            'Segunda pasada de detección de armas con YOLOv8 en alta resolución (imgsz=1600) para cuchillos y armas de fuego.',
            'Implementación de track_vehicles(): seguimiento con historial de posiciones y cálculo de velocidad frame a frame.',
            'Implementación de detect_sudden_deceleration(): umbral configurable de desaceleración brusca (0.7 = pérdida del 70% de velocidad).',
            'Implementación de detect_vehicle_damage(): región de interés del vehículo enviada al modelo detector_de_auto_con_dano.pt para confirmar daños físicos.',
            'Cambio de video_source a None por defecto para compatibilidad frame-by-frame con la UI.',
        ],
        'resultados': 'Pipeline de robos prácticamente completo. Módulo de choques con tracking y análisis de colisiones funcional.',
    },
    {
        'titulo': 'SEMANA 9 — 12 al 16 de mayo de 2026',
        'actividades_plan': 'Detección de robos menores · Detección de choques de autos',
        'horas': '15 horas',
        'actividad_principal': 'Finalización del módulo de robos, integración del módulo de choques y reorganización de alarmas.',
        'acciones': [
            'Optimización del pipeline de robos: reducción de llamadas redundantes a YOLO mediante filtrado unificado de clases. Ajuste de hiperparámetros tras pruebas (umbral celulares: 0.02, umbral personas: 0.50).',
            'Validación del módulo de robos con videos de prueba de tiendas comerciales; precisión estimada: 73%.',
            'Implementación de analyze_accident(): scoring trifactorial — proximidad de vehículos (30%), desaceleración brusca (40%), daño en carrocería (30%); umbral de detección: probabilidad > 0.70.',
            'Implementación de log_accident() y generate_alert(): guardado de imagen del frame e impresión de alerta con timestamp.',
            'Actualización de config.py: ACCIDENT_DETECTION_MODEL apunta a modulos/choques/best.pt; agregado FIRE_DETECTION_MODEL a modulos/incendio/best.pt.',
            'Reorganización del subsistema de alarmas: creación de modulos/alarma/ con telegram_notifier.py, whatsapp_notifier.py y gmail_notifier.py. Eliminación de scripts redundantes whatsapp_selenium.py y whatsapp_simple.py.',
        ],
        'resultados': 'Módulo de robos completo e integrado. Módulo de choques con lógica completa. Subsistema de alarmas reorganizado como paquete Python.',
    },
    {
        'titulo': 'SEMANA 10 — 19 al 22 de mayo de 2026',
        'actividades_plan': 'Detección de choques de autos',
        'horas': '12 horas',
        'actividad_principal': 'Finalización, validación e integración completa del módulo de detección de choques.',
        'acciones': [
            'Integración completa de accident_detection.py en ui.py: el modo "Detección de Accidentes" invoca analyze_accident(frame) en tiempo real y registra incidentes en incidents.db.',
            'Prueba con video de tráfico real: detección correcta de acercamiento brusco con probabilidad > 0.85.',
            'Prueba con video de accidente simulado: daños en carrocería confirmados por detector_de_auto_con_dano.pt con confianza > 0.65.',
            'Correcciones: ajuste de max_distance de 100 a 80 px, limitación del historial a los últimos 30 frames y ajuste de deceleration_threshold de 0.70 a 0.65.',
            'Soporte técnico institucional: diagnóstico y solución de problema de red en el área de seguridad ciudadana; configuración de acceso remoto en dos estaciones de trabajo.',
        ],
        'resultados': 'Módulo de detección de choques completamente funcional, integrado y validado. El sistema acumula tres módulos operativos: incendios, robos menores y choques de autos.',
    },
]

for s in semanas:
    titulo(s['titulo'], level=2, color=AZUL_MEDIO)
    tabla_simple(
        ['Actividades del plan', 'Horas', 'Actividad principal'],
        [[s['actividades_plan'], s['horas'], s['actividad_principal']]],
        col_widths=[5.5, 2.5, 7.5]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Acciones realizadas:')
    r.bold = True
    r.font.size = Pt(11)
    for acc in s['acciones']:
        bullet(acc)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('Resultados: ')
    r2.bold = True
    r2.font.size = Pt(11)
    r2b = p2.add_run(s['resultados'])
    r2b.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  4. ANÁLISIS FODA
# ══════════════════════════════════════════════════════════════════════════════
titulo('4. ANÁLISIS FODA DE LAS PRÁCTICAS')
parrafo(
    'A continuación se presenta la evaluación diagnóstica del proceso de prácticas, identificando factores '
    'internos (Fortalezas y Debilidades) y factores externos (Oportunidades y Amenazas).'
)

titulo('4.1. Fortalezas', level=2)
fortalezas = [
    ('F1', 'Sólidos conocimientos previos en Python, programación orientada a objetos y fundamentos de machine learning adquiridos durante la carrera.'),
    ('F2', 'Capacidad de autoaprendizaje rápido: dominio de nuevas librerías (Ultralytics YOLOv8, PySide6, OpenCV) en tiempo reducido.'),
    ('F3', 'Habilidad para integrar múltiples tecnologías (CUDA, PyTorch, RTSP, SQLite) en un sistema cohesivo y funcional.'),
    ('F4', 'Diseño modular del sistema que facilita el mantenimiento, las pruebas independientes y la incorporación de nuevos módulos.'),
    ('F5', 'Aplicación efectiva de buenas prácticas de ingeniería de software: separación de responsabilidades, configuración centralizada y manejo robusto de excepciones.'),
]
tabla_simple(['#', 'Fortaleza'], fortalezas, col_widths=[1.5, 14])

titulo('4.2. Debilidades', level=2)
debilidades = [
    ('D1', 'Experiencia limitada en el ajuste fino (fine-tuning) de modelos de detección con datasets propios, lo que obligó a depender de modelos preentrenados con datasets genéricos.'),
    ('D2', 'Conocimiento inicial insuficiente en protocolos RTSP y gestión de streams de video en tiempo real, generando retrasos en la semana de captura de video.'),
    ('D3', 'Dificultad para optimizar el rendimiento en GPU de baja potencia (2 GB VRAM), requiriendo reducción de resolución de inferencia y ajuste manual de parámetros.'),
    ('D4', 'Limitaciones en el tiempo disponible (3 horas/día) que obligaron a priorizar funcionalidades críticas sobre características secundarias.'),
]
tabla_simple(['#', 'Debilidad'], debilidades, col_widths=[1.5, 14])

titulo('4.3. Oportunidades', level=2)
oportunidades = [
    ('O1', 'Creciente disponibilidad de modelos preentrenados de código abierto (YOLOv8, RT-DETR) que reducen el tiempo de desarrollo.'),
    ('O2', 'Interés institucional genuino de la municipalidad en adoptar tecnología de IA, asegurando continuidad y escalabilidad del proyecto.'),
    ('O3', 'Posibilidad de extender el sistema hacia una plataforma web con acceso remoto multiusuario para otras instituciones de la región.'),
    ('O4', 'Disponibilidad de datasets públicos de alta calidad para incendios (Fire-8) y accidentes que pueden usarse para reentrenamiento especializado.'),
    ('O5', 'Potencial de publicación académica de los resultados en revistas de ingeniería o ponencias universitarias.'),
]
tabla_simple(['#', 'Oportunidad'], oportunidades, col_widths=[1.5, 14])

titulo('4.4. Amenazas', level=2)
amenazas = [
    ('A1', 'Infraestructura de red limitada: ancho de banda insuficiente para streaming de múltiples cámaras a alta resolución simultáneamente.'),
    ('A2', 'Riesgo de cortes de energía y falta de UPS en el área de servidores, pudiendo interrumpir el sistema en momentos críticos.'),
    ('A3', 'Variación en condiciones de iluminación (noche, contraluz, niebla altiplánica) que reduce la efectividad de detectores entrenados con imágenes diurnas.'),
    ('A4', 'Posible obsolescencia del hardware disponible (Quadro P1000, 2016) que limitaría la ejecución de modelos más precisos en el futuro.'),
]
tabla_simple(['#', 'Amenaza'], amenazas, col_widths=[1.5, 14])

# ══════════════════════════════════════════════════════════════════════════════
#  5. SUGERENCIAS
# ══════════════════════════════════════════════════════════════════════════════
titulo('5. SUGERENCIAS')

titulo('5.1. Para la Escuela Profesional de Ingeniería de Software y Sistemas — UNAJ', level=2)
bullet('Incorporar cursos de visión por computadora e inteligencia artificial aplicada en el plan de estudios, con énfasis práctico en frameworks como PyTorch y Ultralytics YOLOv8.')
bullet('Fortalecer el curso de Ingeniería de Software con proyectos de integración modular real, cubriendo el ciclo completo: análisis de requerimientos → diseño → implementación → pruebas → despliegue.')
bullet('Establecer convenios de prácticas con instituciones públicas (municipalidades, hospitales) para que los proyectos de los practicantes tengan impacto social directo y medible.')
bullet('Proveer acceso a laboratorios con GPUs (al menos GTX 1660 o equivalente) para que los estudiantes puedan entrenar y validar modelos de deep learning durante los cursos.')

titulo('5.2. Para la Municipalidad Distrital de Caracoto', level=2)
bullet('Actualizar la infraestructura de red del área de Seguridad Ciudadana con un switch gestionable dedicado para el sistema de cámaras y un mínimo de 20 Mbps por cámara.')
bullet('Instalar un sistema de alimentación ininterrumpida (UPS) en el servidor que aloja el sistema CCTV AI PRO para evitar pérdida de datos en cortes de luz.')
bullet('Considerar la adquisición de hardware más potente (GPU NVIDIA RTX 3060, 8 GB VRAM mínimo) que permita ejecutar los módulos de detección en todas las cámaras simultáneamente.')
bullet('Recolectar y etiquetar imágenes locales de situaciones de riesgo reales para generar un dataset propio y reentrenar los modelos con mayor precisión en el contexto de Caracoto.')
bullet('Designar un operador capacitado para el uso y mantenimiento del sistema, con formación básica en Python y configuración de cámaras IP.')
bullet('Activar el sistema de notificaciones (Telegram y Gmail en modulos/alarma/) configurando credenciales reales para que las alertas lleguen automáticamente al personal del serenazgo.')

# ══════════════════════════════════════════════════════════════════════════════
#  6. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
titulo('6. CONCLUSIONES')

conclusiones = [
    ('Primera.', 'Se implementaron satisfactoriamente tres módulos de detección de incidentes críticos — incendios, robos menores y choques de vehículos — integrándolos en el sistema cohesivo CCTV AI PRO en tiempo real, cumpliendo el objetivo principal de las prácticas preprofesionales dentro del cronograma establecido.'),
    ('Segunda.', 'El módulo de detección de incendios resultó ser el más robusto y confiable, gracias a la combinación de un detector YOLOv8 entrenado sobre el dataset Fire-8 con un análisis de respaldo por color HSV. Esta detección híbrida logró una tasa de detección estimada del 87% en condiciones de iluminación normal, con bajo número de falsos positivos.'),
    ('Tercera.', 'La implementación de la arquitectura modular con separación clara de responsabilidades facilitó el desarrollo paralelo de componentes, simplificó las pruebas unitarias y garantizó la mantenibilidad del código a largo plazo, siguiendo los principios de ingeniería de software de Pressman & Maxim (2014).'),
    ('Cuarta.', 'El módulo de detección de choques de autos presentó el mayor desafío técnico, al requerir seguimiento simultáneo de vehículos, cálculo de velocidades relativas y confirmación de daños mediante un segundo modelo de IA. La refactorización del constructor para soporte frame-by-frame fue clave para su integración correcta en la UI.'),
    ('Quinta.', 'Las prácticas permitieron consolidar competencias técnicas avanzadas en ingeniería de software aplicada: integración de modelos de deep learning en aplicaciones de escritorio, gestión de hilos para procesamiento de video en tiempo real, diseño de bases de datos embebidas y construcción de interfaces gráficas responsivas con PySide6.'),
]

for label, text in conclusiones:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

# ── Firma ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
centrado('Juliaca, 15 de julio del 2026', espacio_antes=20)
doc.add_paragraph()
centrado('_________________________________')
centrado('Yords Williams Ccalla Mamani', bold=True)
centrado('Practicante — Municipalidad Distrital de Caracoto')
centrado('Escuela Profesional de Ingeniería de Software y Sistemas — UNAJ')

# ── Guardar ───────────────────────────────────────────────────────────────────
output = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'INFORME_FINAL_PRACTICAS.docx')
doc.save(output)
print(f'Guardado: {output}')
