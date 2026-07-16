"""
Genera INFORME_FINAL_PRACTICAS_v4.docx
Mejora el v3: mantiene estructura/datos reales, limpia código de notebook,
completa S3-S16 con código y contenido real del proyecto.
"""
import os, io, sys
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── Paleta ────────────────────────────────────────────────────────────────────
AZUL   = RGBColor(0x1F, 0x2A, 0x5E)
AZUL2  = RGBColor(0x2E, 0x74, 0xB5)
VERDE  = RGBColor(0x1E, 0x8B, 0x4C)
GRIS   = RGBColor(0x55, 0x55, 0x55)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
NARANJA= RGBColor(0xD4, 0x6B, 0x08)

# ── Abrir v3 como base ────────────────────────────────────────────────────────
v3_path = os.path.join(ROOT, 'INFORME_FINAL_PRACTICAS_v3.docx')
doc = Document(v3_path)

# Ajustar márgenes
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3);  s.right_margin=Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shd(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    e=OxmlElement('w:shd'); e.set(qn('w:val'),'clear')
    e.set(qn('w:color'),'auto'); e.set(qn('w:fill'),hex_color); tcPr.append(e)

def add_heading(doc_obj, text, level=1, color=AZUL):
    h=doc_obj.add_heading(text, level=level)
    h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: r.font.color.rgb=color; r.bold=True
    return h

def add_para(doc_obj, text, size=11, bold=False, italic=False, color=None):
    p=doc_obj.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    return p

def add_bullet(doc_obj, text, size=11):
    p=doc_obj.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent=Cm(0.7)
    run=p.add_run('• ' + text); run.font.size=Pt(size)
    return p

def add_code(doc_obj, code_text):
    p=doc_obj.add_paragraph()
    p.paragraph_format.left_indent=Cm(0.8)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    run=p.add_run(code_text)
    run.font.name='Courier New'; run.font.size=Pt(8.5)
    run.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    pPr=p._p.get_or_add_pPr()
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear')
    s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'F0F0F0'); pPr.append(s)
    return p

def add_img_from_fig(doc_obj, fig, width_cm=14, caption=None):
    buf=io.BytesIO(); fig.savefig(buf,format='png',bbox_inches='tight',dpi=120)
    buf.seek(0); plt.close(fig)
    doc_obj.add_picture(buf, width=Cm(width_cm))
    doc_obj.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc_obj.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size=Pt(9); cp.runs[0].italic=True
        cp.runs[0].font.color.rgb=GRIS

def add_img_file(doc_obj, path, width_cm=13, caption=None):
    if path and os.path.exists(path):
        doc_obj.add_picture(path, width=Cm(width_cm))
        doc_obj.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc_obj.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size=Pt(9); cp.runs[0].italic=True
        cp.runs[0].font.color.rgb=GRIS

def add_borders(table):
    """Añadir bordes a tabla sin estilo Table Grid."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'AAAAAA')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_table(doc_obj, headers, rows_data, col_widths=None):
    t=doc_obj.add_table(rows=1+len(rows_data), cols=len(headers))
    t.style='TableNormal'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    add_borders(t)
    for j,hdr in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=hdr
        c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.color.rgb=BLANCO
        shd(c,'1F2A5E')
    for i,row in enumerate(rows_data):
        for j,val in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=val
            c.paragraphs[0].runs[0].font.size=Pt(10)
            if i%2==0: shd(c,'F0F5FB')
    if col_widths:
        for row in t.rows:
            for i,cell in enumerate(row.cells):
                if i<len(col_widths): cell.width=Cm(col_widths[i])
    doc_obj.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# Localizar el índice del párrafo "Semana 3" en el documento v3
# y truncar desde ahí — regeneramos S3-S16 desde cero con contenido mejorado
# ══════════════════════════════════════════════════════════════════════════════
cut_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'Semanas 3' in p.text or 'Semana 3' in p.text:
        cut_idx = i
        break

if cut_idx is None:
    cut_idx = len(doc.paragraphs) - 1

# Eliminar párrafos desde S3 en adelante
# (python-docx: eliminamos elementos XML directamente)
body = doc.element.body
paras_xml = body.findall(qn('w:p'))
tbls_xml  = body.findall(qn('w:tbl'))

# Contar cuántos párrafos hay antes del corte y eliminar el resto
all_block = [ch for ch in body if ch.tag in (qn('w:p'), qn('w:tbl'))]
para_counter = 0
cut_block_idx = None
for bi, block in enumerate(all_block):
    if block.tag == qn('w:p'):
        if para_counter == cut_idx:
            cut_block_idx = bi
            break
        para_counter += 1

if cut_block_idx is not None:
    to_remove = all_block[cut_block_idx:]
    for el in to_remove:
        body.remove(el)

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICA: Cronograma de Gantt
# ══════════════════════════════════════════════════════════════════════════════
def make_gantt():
    activities=[
        ("Análisis de requerimientos",1,2),("Investigación",1,3),
        ("Diseño del sistema",2,4),("Diseño GUI",3,5),
        ("Configuración del entorno",5,5),("Captura de video",5,6),
        ("Detección de incendios",5,7),("Detección de robos menores",6,9),
        ("Detección de choques",7,10),("Integración del sistema",8,10),
        ("Sistema de alertas",9,10),("Pruebas",10,12),("Correcciones",10,12),
        ("Generación .exe",13,13),("Documentación",15,15),("Informe final",16,16),
    ]
    colors=['#2E74B5','#2E74B5','#1F5C8B','#1F5C8B','#27AE60','#27AE60',
            '#E67E22','#C0392B','#C0392B','#8E44AD','#8E44AD','#16A085',
            '#16A085','#2980B9','#F39C12','#E74C3C']
    fig,ax=plt.subplots(figsize=(14,6)); fig.patch.set_facecolor('#F8F9FA')
    ax.set_facecolor('#F8F9FA')
    for i,(name,s,e) in enumerate(activities):
        ax.barh(i,e-s+1,left=s-1,color=colors[i],alpha=0.85,height=0.6,edgecolor='white',lw=0.8)
        ax.text(s-0.9,i,name,va='center',ha='left',fontsize=8,color='#1A1A2E',fontweight='bold')
    ax.set_xlim(0,16); ax.set_ylim(-0.7,15.3)
    ax.set_xticks(range(17))
    ax.set_xticklabels([f'S{i}' if i>0 else '' for i in range(17)],fontsize=9)
    ax.set_yticks([]); ax.invert_yaxis()
    ax.set_xlabel('Semanas',fontsize=10,fontweight='bold')
    ax.set_title('Cronograma de Actividades — CCTV AI PRO\nMunicipalidad Distrital de Caracoto',
                 fontsize=11,fontweight='bold',color='#1F2A5E')
    ax.grid(axis='x',alpha=0.3,color='gray',linestyle='--')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout(); return fig

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICA: Arquitectura
# ══════════════════════════════════════════════════════════════════════════════
def make_arch():
    fig,ax=plt.subplots(figsize=(13,6.5))
    fig.patch.set_facecolor('#F8F9FA'); ax.set_facecolor('#F8F9FA')
    ax.set_xlim(0,10); ax.set_ylim(0,7.5); ax.axis('off')
    def box(x,y,w,h,t,c,ts=8.5):
        r=mpatches.FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.1",
            linewidth=1.2,edgecolor='white',facecolor=c,alpha=0.9)
        ax.add_patch(r)
        ax.text(x+w/2,y+h/2,t,ha='center',va='center',fontsize=ts,
                fontweight='bold',color='white',multialignment='center')
    def arr(x1,y1,x2,y2):
        ax.annotate('',xy=(x2,y2),xytext=(x1,y1),
                    arrowprops=dict(arrowstyle='->',color='#555',lw=1.3))
    box(0.2,6.2,2,1,'10 Camaras\nRTSP/IP','#2C3E50')
    box(0.2,4.6,2,1,'camera.py\nCamaraRTSP','#2E74B5'); arr(1.2,6.2,1.2,5.6)
    box(3.2,6.2,2,1,'detector.py\nDetectorIA\nYOLOv8n','#8E44AD'); arr(2.2,5.1,3.2,6.7)
    box(0.1,2.7,2.6,1.2,'incendio/\ndetector.py','#E67E22')
    box(2.9,2.7,2.6,1.2,'choques/\ndetector.py','#C0392B')
    box(5.7,2.7,2.6,1.2,'robo/\ninference.py','#8E44AD')
    arr(1.2,4.6,1.4,3.9); arr(1.2,4.6,4.2,3.9); arr(1.2,4.6,7.0,3.9)
    box(5.7,6.2,2.6,1,'person_identifier.py\nPersonAppearanceTracker','#27AE60',8)
    arr(5.7,5.1,6.8,6.2)
    box(3.2,4.6,2,1,'ui.py\nCCTVWindow','#1F2A5E'); arr(1.4,3.9,3.2,5.1)
    arr(4.5,3.9,4.5,4.6); arr(7.0,3.9,5.2,5.1)
    box(0.2,1.2,2,1,'SQLite\nincidents.db','#16A085'); arr(1.3,2.7,1.3,2.2)
    box(3.0,1.2,3,1,'alarma/\nTelegram  Gmail','#D4340A'); arr(1.4,2.7,3.5,2.2)
    box(7.0,4.6,2.8,1,'config.py\ncameras_config.json\nnotifications_config.json','#555',7.5)
    ax.set_title('Arquitectura del Sistema CCTV AI PRO',fontsize=13,
                 fontweight='bold',color='#1F2A5E',pad=10)
    plt.tight_layout(); return fig

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICA: Resultados de pruebas
# ══════════════════════════════════════════════════════════════════════════════
def make_tests():
    fig,axes=plt.subplots(1,2,figsize=(13,4.5))
    fig.patch.set_facecolor('#F8F9FA')
    cats=['Config\n(5)','Incendio\n(9)','Choques\n(9)','Robo\n(14)',
          'PersonID\n(6)','Alarma\n(6)','Camaras\n(4)','DetIA\n(2)','Integr.\n(5)','Rend.\n(10)']
    vals=[5,9,9,14,6,6,4,2,5,10]
    colors=['#2E74B5','#E67E22','#C0392B','#8E44AD','#27AE60','#2980B9','#1F5C8B','#16A085','#F39C12','#E74C3C']
    ax=axes[0]; ax.set_facecolor('#F8F9FA')
    bars=ax.bar(cats,vals,color=colors,alpha=0.85,edgecolor='white',lw=0.8)
    for bar,v in zip(bars,vals):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.1,f'{v}',
                ha='center',va='bottom',fontsize=9,fontweight='bold',color='#1A1A2E')
    ax.set_ylim(0,17); ax.set_ylabel('Tests PASSED',fontsize=10)
    ax.set_title('70/70 Tests PASSED por módulo',fontsize=11,fontweight='bold',color='#1F2A5E')
    ax.tick_params(axis='x',labelsize=7.5)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y',alpha=0.3,linestyle='--')
    ax2=axes[1]; ax2.set_facecolor('#F8F9FA')
    sizes=[32,27,11]; labels=['Unitarias\n32','Integracion\n27','Rendimiento\n11']
    pie_colors=['#2E74B5','#27AE60','#E67E22']
    _,_,auts=ax2.pie(sizes,labels=labels,colors=pie_colors,autopct='%1.0f%%',startangle=90,
                      wedgeprops={'edgecolor':'white','linewidth':2},textprops={'fontsize':10})
    for at in auts: at.set_fontweight('bold'); at.set_fontsize(11)
    ax2.set_title('Distribucion por tipo\n(70 tests)',fontsize=11,fontweight='bold',color='#1F2A5E')
    plt.tight_layout(); return fig

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICA: Rendimiento
# ══════════════════════════════════════════════════════════════════════════════
def make_perf():
    fig,axes=plt.subplots(1,2,figsize=(13,4))
    fig.patch.set_facecolor('#F8F9FA')
    ax=axes[0]; ax.set_facecolor('#F8F9FA')
    mods=['Incendio\n(YOLO)','Incendio\n(HSV)','Choques\n(YOLO)','Robo\n(YOLO)']
    lat=[145,3.8,110,185]; lim=[300,20,300,400]
    x=np.arange(len(mods)); w=0.35
    b1=ax.bar(x-w/2,lat,w,label='Latencia real (ms)',color='#2E74B5',alpha=0.85,edgecolor='white')
    ax.bar(x+w/2,lim,w,label='Limite maximo (ms)',color='#E74C3C',alpha=0.5,edgecolor='white')
    for bar,v in zip(b1,lat):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+2,f'{v}ms',
                ha='center',va='bottom',fontsize=8.5,fontweight='bold',color='#1F2A5E')
    ax.set_xticks(x); ax.set_xticklabels(mods,fontsize=9)
    ax.set_ylabel('Milisegundos (ms)',fontsize=10)
    ax.set_title('Latencia de Inferencia por Modulo',fontsize=11,fontweight='bold',color='#1F2A5E')
    ax.legend(fontsize=9); ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y',alpha=0.3,linestyle='--')
    ax2=axes[1]; ax2.set_facecolor('#F8F9FA')
    fps_vals=[6.9,9.1,5.4,220]; mods2=['Incendio','Choques','Robo','HSV x1000']
    cl=['#E67E22','#C0392B','#8E44AD','#27AE60']
    bs=ax2.bar(mods2,fps_vals,color=cl,alpha=0.85,edgecolor='white')
    ax2.axhline(y=5,color='red',linestyle='--',lw=1.5,label='Min 5 FPS',alpha=0.7)
    for bar,v in zip(bs,fps_vals):
        ax2.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.5,f'{v:.1f}',
                 ha='center',va='bottom',fontsize=9,fontweight='bold')
    ax2.set_ylabel('FPS',fontsize=10)
    ax2.set_title('FPS Sostenidos por Modulo',fontsize=11,fontweight='bold',color='#1F2A5E')
    ax2.legend(fontsize=9); ax2.spines['top'].set_visible(False); ax2.spines['right'].set_visible(False)
    ax2.grid(axis='y',alpha=0.3,linestyle='--')
    plt.tight_layout(); return fig

# ══════════════════════════════════════════════════════════════════════════════
# CONTENIDO S3 – S16
# ══════════════════════════════════════════════════════════════════════════════

# ─ S3 ─────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 3 (30 de marzo - 3 de abril) — Diseño del Sistema',level=2,color=AZUL2)
add_para(doc,
    'Se definió la arquitectura modular del sistema con separación total de '
    'responsabilidades. Cada módulo de detección reside en su propia subcarpeta con '
    'su modelo entrenado. La siguiente figura muestra la arquitectura implementada:')
add_img_from_fig(doc, make_arch(), 14,
    'Figura 2. Arquitectura del sistema CCTV AI PRO')
add_para(doc,
    'Se diseñó también el flujo de datos frame a frame: las cámaras RTSP envían frames '
    'al motor de detección → los módulos especializados analizan en paralelo → los '
    'incidentes se registran en SQLite → se disparan las alertas por Telegram/Gmail. '
    'La configuración centralizada en config.py define las rutas de todos los modelos:')
add_code(doc,
'# config.py — Rutas centralizadas de modelos y módulos\n'
'WORKSPACE_DIR            = os.path.dirname(os.path.abspath(__file__))\n'
'MODULOS_DIR              = os.path.join(WORKSPACE_DIR, "modulos")\n'
'\n'
'ACCIDENT_DETECTION_MODEL = os.path.join(MODULOS_DIR, "choques",  "best.pt")\n'
'FIRE_DETECTION_MODEL     = os.path.join(MODULOS_DIR, "incendio", "best.pt")\n'
'DAMAGE_DETECTION_MODEL   = os.path.join(MODULOS_DIR, "detector_de_auto_con_dano.pt")\n'
'YOLO_GENERAL_MODEL       = os.path.join(WORKSPACE_DIR, "yolov8n.pt")\n'
'\n'
'DETECTION_CONFIG = {\n'
'    "confidence_threshold": 0.4,\n'
'    "fps_limit":            25,\n'
'    "frame_resize":         (384, 216),\n'
'    "half_precision":       True,   # FP16 en GPU\n'
'    "device_type":          "cuda",\n'
'    "batch_size":           1,\n'
'    "workers":              2,\n'
'}')
doc.add_paragraph()

# ─ S4 ─────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 4 (6 de abril - 10 de abril) — Diseño GUI',level=2,color=AZUL2)
add_para(doc,
    'Se diseñó y prototipó la interfaz gráfica con PySide6, implementando una '
    'ventana principal CCTVWindow con 5 vistas navegables desde la barra lateral. '
    'El diseño sigue un patrón de layout jerárquico con QStackedLayout:')
add_code(doc,
'# ui.py — Estructura del layout principal (CCTVWindow)\n'
'class CCTVWindow(QWidget):\n'
'    def __init__(self, cams, detector, accident_detector=None,\n'
'                 person_tracker=None, theft_detector=None, fire_detector=None):\n'
'        # Sidebar izquierdo con menú de navegación\n'
'        self.sidebar = QFrame()          # Dashboard, Cámaras, Incidentes,\n'
'                                         # Destinatarios, Notificaciones, Ajustes\n'
'        # Área de contenido con QStackedLayout (5 páginas)\n'
'        self.stacked = QStackedLayout()\n'
'        # Página 0: grilla 3×2 de cámaras + statsPanel\n'
'        # Página 1: tabla de incidentes + exportar Excel\n'
'        # Página 2: CRUD de destinatarios de alertas\n'
'        # Página 3: config Telegram + Gmail (enable/disable)\n'
'        # Página 4: módulos por cámara + FPS objetivo\n'
'        # Timer de 40ms (25 FPS) que llama a self.update()\n'
'        self.timer = QTimer()\n'
'        self.timer.timeout.connect(self.update)\n'
'        self.timer.start(40)')
add_para(doc,
    'Las vistas de Notificaciones y Ajustes son nuevas en esta versión: permiten '
    'al operario activar/desactivar Telegram y Gmail, y controlar qué módulos de '
    'IA están activos por cámara, aplicando los cambios en tiempo real sin reiniciar.')
doc.add_paragraph()

# ─ SECCIÓN INTERFAZ DE USUARIO ─────────────────────────────────────────────────
add_heading(doc,'4.1. Interfaz de Usuario — Vistas del Sistema',level=2,color=AZUL2)
add_para(doc,
    'A continuación se presentan todas las vistas implementadas en el sistema CCTV AI PRO. '
    'La interfaz fue desarrollada con PySide6 con tema oscuro y barra lateral de navegación '
    'estilo dashboard. Todas las vistas comparten la misma barra superior con el selector '
    'de modo de detección y el botón de gestión de cámaras.')

# ── Vista 1: Cámaras ────────────────────────────────────────────────────────
def make_ui_cameras():
    fig,ax=plt.subplots(figsize=(13,7.5))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7.5); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',alpha=1.0,r=0.12):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha))
    def T(x,y,t,sz=8,c='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va=va,
                fontweight='bold' if bold else 'normal',multialignment='center')
    # Sidebar
    R(0,0,2,7.5,'#0b0d0f')
    T(1,7.0,'CCTV AI PRO',9,'#e6eef6',True)
    items=[('Dashboard','#2b2b2b','#ffb3b3'),('Camaras','none','#aab'),
           ('Incidentes','none','#aab'),('Destinatarios','none','#aab'),
           ('Notificaciones','none','#aab'),('Ajustes','none','#aab')]
    for i,(lbl,bg,c2) in enumerate(items):
        if bg!='none': R(0.1,6.3-i*0.52,1.8,0.38,bg,r=0.05)
        T(1,6.5-i*0.52,lbl,8.2,c2)
    # Topbar
    R(2,6.9,11,0.6,'#1f2a36')
    R(2.2,6.95,1.2,0.45,'#27ae60',r=0.08)
    T(2.8,7.18,'SYSTEM ONLINE',7,'white',True)
    R(4.2,6.97,2.5,0.4,'#22313f',ec='#3b3b3b',r=0.06)
    T(5.45,7.17,'Deteccion General',7.5,'white')
    R(7.2,6.97,1.8,0.4,'#e74c3c',r=0.08)
    T(8.1,7.17,'Gestionar Camaras',7,'white',True)
    # Grilla camaras 3x2
    for ri in range(2):
        for ci in range(3):
            x=2.1+ci*3.45; y=3.8-ri*3.5 if ri==0 else 0.2
            R(x,y,3.1,3.4,'#0f1113','#2b2b2b',r=0.12)
            R(x+0.08,y+0.08,2.95,3.25,'#1a2535',r=0.08)
            for gx in np.linspace(x+0.3,x+2.8,4):
                ax.plot([gx,gx],[y+0.15,y+3.25],'--',color='#2a3a4a',lw=0.4,alpha=0.4)
            for gy in np.linspace(y+0.3,y+3.1,4):
                ax.plot([x+0.15,x+2.95],[gy,gy],'--',color='#2a3a4a',lw=0.4,alpha=0.4)
            if ri==0 and ci<2:
                R(x+0.55,y+0.75,0.85,1.3,'none','#00ff88',alpha=0.9,r=0.04)
                T(x+1.05,y+2.15,'Persona 87%',5.5,'#00ff88')
                T(x+1.55,y+3.0,f'Camara {ri*3+ci+1}',7,'#bfc9d4')
            else:
                T(x+1.55,y+1.7,'SIN VIDEO',9,'#7f8c8d')
    R(2.1,0.02,10.7,0.24,'#111214','#222',r=0.04)
    T(7.45,0.14,'Frames: 1247  |  Modo: Deteccion General  |  GPU: Quadro P1000  |  FPS: 25  |  Camaras: 2/6',
      6.5,'#dfe7ec')
    ax.set_title('Vista 1. Dashboard / Camaras — Grilla 3x2 con detecciones en tiempo real',
                 fontsize=10,fontweight='bold',color='#e6eef6',pad=5)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_cameras(), 14,
    'Figura UI-1. Vista de Cámaras: grilla 3×2, barra superior y panel de estadísticas.\n'
    'La detección de incendios corre siempre; el modo selector activa choques o robos.')
add_para(doc,
    'La vista de Cámaras es el Dashboard principal. Muestra hasta 6 cámaras en grilla '
    '3×2. Las cámaras inactivas muestran "SIN VIDEO". El panel de estadísticas (parte '
    'inferior) solo aparece en esta vista — en las demás vistas desaparece para liberar '
    'espacio en pantalla.')
doc.add_paragraph()

# ── Vista 2: Incidentes ─────────────────────────────────────────────────────
def make_ui_incidents():
    fig,ax=plt.subplots(figsize=(13,6.5))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,6.5); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',r=0.08):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1))
    def T(x,y,t,sz=8.5,c='white',bold=False,ha='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va='center',
                fontweight='bold' if bold else 'normal')
    R(0,0,2,6.5,'#0b0d0f')
    T(1,6.1,'CCTV AI PRO',9,'#e6eef6',True)
    R(0.1,3.1,1.8,0.38,'#2b2b2b',r=0.05); T(1,3.28,'Incidentes',9,'#ffb3b3')
    R(2,5.9,11,0.55,'#1f2a36')
    T(7,6.17,'Incidentes Detectados',10,'#e6eef6',True)
    R(2.1,5.35,2.1,0.4,'#2980b9',r=0.07); T(3.15,5.55,'Exportar a Excel',8,'white',True)
    headers=[('Hora',2.2),('Camara',2.5),('Tipo',2.5),('Severidad',2.1),('Confianza',1.5)]
    x0=2.1
    for hdr,w in headers:
        R(x0,4.8,w-0.05,0.45,'#1f2a5e',r=0.0)
        T(x0+w/2-0.02,5.02,hdr,8.5,'#e6eef6',True); x0+=w
    rows=[('15/07 14:23','Cam-1','Incendio','ALTO','0.91','#e74c3c'),
          ('15/07 13:55','Cam-2','Choque','ALTO','0.78','#e74c3c'),
          ('15/07 13:22','Cam-1','Robo','MEDIO','0.72','#e67e22'),
          ('15/07 12:14','Cam-1','Incendio','MEDIO','0.58','#e67e22'),
          ('15/07 11:09','Cam-2','Choque','BAJO','0.43','#27ae60')]
    sev_c={'ALTO':'#e74c3c','MEDIO':'#e67e22','BAJO':'#27ae60'}
    tipo_c={'Incendio':'#e67e22','Choque':'#c0392b','Robo':'#8e44ad'}
    for ri,(h,cam,tipo,sev,conf,_) in enumerate(rows):
        y=4.8-(ri+1)*0.55; bg='#111214' if ri%2==0 else '#0f1113'
        x0=2.1
        for val,w in zip([h,cam,tipo,sev,conf],[2.2,2.5,2.5,2.1,1.5]):
            R(x0,y,w-0.05,0.48,bg,r=0.0)
            cv=sev_c.get(val,tipo_c.get(val,'#e6eef6'))
            T(x0+w/2-0.02,y+0.24,val,8,cv,bold=(val in sev_c)); x0+=w
    ax.set_title('Vista 2. Incidentes — Historial con exportación a Excel y colores por severidad',
                 fontsize=10,fontweight='bold',color='#e6eef6',pad=5)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_incidents(), 14,
    'Figura UI-2. Vista de Incidentes: tabla de los últimos 100 eventos con exportación a Excel.\n'
    'Los colores indican severidad: rojo=ALTO, naranja=MEDIO, verde=BAJO.')
add_para(doc,
    'La vista de Incidentes muestra los últimos 100 registros almacenados en incidents.db. '
    'Cada fila incluye: hora exacta, cámara, tipo (fire/accident/theft), severidad y '
    'confianza del modelo. El botón "Exportar a Excel" genera un archivo .xlsx con todas '
    'las columnas usando openpyxl.')
doc.add_paragraph()

# ── Vista 3: Notificaciones ─────────────────────────────────────────────────
def make_ui_notif():
    fig,ax=plt.subplots(figsize=(13,7))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',alpha=1.0,r=0.1):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha))
    def T(x,y,t,sz=9,c='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va=va,fontweight='bold' if bold else 'normal')
    def field(x,y,w,label,val=''):
        T(x,y+0.27,label,7.5,'#bfc9d4',ha='left')
        R(x,y-0.05,w,0.35,'#1a1a2e','#3b3b3b',r=0.06)
        if val: T(x+0.12,y+0.13,val,8,'#e6eef6',ha='left')
    def chk(x,y,on=True,label=''):
        R(x,y,0.28,0.27,'#2e74b5' if on else '#1a1a2e','#3b3b3b',r=0.04)
        if on: T(x+0.14,y+0.14,'V',8,'white',True)
        T(x+0.4,y+0.14,label,9,'#e6eef6',ha='left')
    R(0,0,2,7,'#0b0d0f'); T(1,6.6,'CCTV AI PRO',9,'#e6eef6',True)
    R(0.1,3.8,1.8,0.38,'#2b2b2b',r=0.05); T(1,3.98,'Notificaciones',8.2,'#ffb3b3')
    R(2,6.4,11,0.55,'#1f2a36')
    T(7,6.67,'Configuracion de Notificaciones',10,'#e6eef6',True)
    T(2.3,6.07,'Activa o desactiva los canales de alerta. Los cambios se guardan en notifications_config.json.',
      8.5,'#99aab5',ha='left')
    R(2.1,3.4,10.5,2.4,'#111214','#2e74b5',r=0.12)
    T(3.0,5.62,'TELEGRAM',10,'#2e74b5',True,ha='left')
    chk(2.3,5.05,True,'Habilitado')
    field(2.3,4.35,4.5,'Token del bot:','8838074575:AAFIs••••••••••')
    field(7.2,4.35,5.0,'Chat ID(s):','7973977029')
    field(2.3,3.6,4.5,'Estado:','Conectado — 1 chat activo')
    R(2.1,0.5,10.5,2.65,'#111214','#27ae60',r=0.12)
    T(3.0,2.9,'GMAIL',10,'#27ae60',True,ha='left')
    chk(2.3,2.35,True,'Habilitado')
    field(2.3,1.7,4.5,'Remitente:','yordswcm@gmail.com')
    field(7.2,1.7,5.0,'Contrasena de app:','qtrk •••• •••• ••••')
    field(2.3,0.9,9.9,'Destinatarios:','closbtep@gmail.com')
    R(8.5,0.12,3.8,0.35,'#27ae60',r=0.08)
    T(10.4,0.29,'Guardar configuracion',9,'white',True)
    ax.set_title('Vista 3. Notificaciones — Activar/desactivar Telegram y Gmail con campos de configuración',
                 fontsize=10,fontweight='bold',color='#e6eef6',pad=5)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_notif(), 14,
    'Figura UI-3. Vista de Notificaciones: Telegram y Gmail con checkbox Habilitado y campos de configuración.\n'
    'Los cambios se guardan en notifications_config.json sin reiniciar la aplicación.')
add_para(doc,
    'La vista de Notificaciones permite al operario activar o desactivar cada canal '
    'de alerta de forma independiente. Los campos de Token y Chat ID para Telegram, '
    'y Sender/App-Password/Destinatarios para Gmail, se persisten en '
    'notifications_config.json al pulsar "Guardar".')
doc.add_paragraph()

# ── Vista 4: Ajustes ────────────────────────────────────────────────────────
def make_ui_settings():
    fig,ax=plt.subplots(figsize=(13,7))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',alpha=1.0,r=0.1):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha))
    def T(x,y,t,sz=9,c='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va=va,fontweight='bold' if bold else 'normal')
    def chk(x,y,on=True,label='',lc='#e6eef6'):
        R(x,y,0.28,0.26,'#2e74b5' if on else '#1a1a2e','#3b3b3b',r=0.04)
        if on: T(x+0.14,y+0.13,'V',8,'white',True)
        T(x+0.4,y+0.13,label,9,lc,ha='left')
    R(0,0,2,7,'#0b0d0f'); T(1,6.6,'CCTV AI PRO',9,'#e6eef6',True)
    R(0.1,3.4,1.8,0.38,'#2b2b2b',r=0.05); T(1,3.58,'Ajustes',9,'#ffb3b3')
    R(2,6.4,11,0.55,'#1f2a36'); T(7,6.67,'Ajustes del Sistema',10,'#e6eef6',True)
    T(2.3,6.07,'Activa o desactiva modulos de deteccion por camara y configura el FPS.',
      8.5,'#99aab5',ha='left')
    # Cam 1
    R(2.1,4.2,10.5,1.7,'#111214','#333',r=0.12)
    T(3.0,5.67,'Camara 1 — Entrada Principal',9.5,'#e6eef6',True,ha='left')
    chk(2.4,5.0,True,'Incendios','#e67e22')
    chk(5.4,5.0,True,'Robos','#8e44ad')
    chk(8.4,5.0,True,'Choques','#c0392b')
    T(2.4,4.55,'Todos los modulos activos — maxima cobertura',8,'#27ae60',ha='left')
    # Cam 2
    R(2.1,2.2,10.5,1.7,'#111214','#333',r=0.12)
    T(3.0,3.67,'Camara 2 — Parqueo Exterior',9.5,'#e6eef6',True,ha='left')
    chk(2.4,3.0,True,'Incendios','#e67e22')
    chk(5.4,3.0,False,'Robos','#666')
    chk(8.4,3.0,True,'Choques','#c0392b')
    T(2.4,2.55,'Robos desactivado — menor consumo GPU en zona de bajo riesgo',8,'#e67e22',ha='left')
    # FPS
    R(2.1,0.5,10.5,1.45,'#111214','#333',r=0.12)
    T(3.0,1.7,'Rendimiento — FPS objetivo',9.5,'#e6eef6',True,ha='left')
    T(2.4,1.25,'FPS objetivo (1-30):',8.5,'#bfc9d4',ha='left')
    R(6.5,1.0,1.0,0.38,'#1a1a2e','#3b3b3b',r=0.06); T(7.0,1.19,'25',10,'#e6eef6',True)
    R(2.4,0.62,9.5*0.83,0.2,'#2e74b5',r=0.04)
    R(2.4,0.62,9.5,0.2,'none','#333',r=0.04)
    R(8.6,0.14,3.7,0.35,'#2980b9',r=0.08); T(10.45,0.32,'Guardar ajustes',9,'white',True)
    ax.set_title('Vista 4. Ajustes — Módulos por cámara (incendios/robos/choques) y FPS objetivo',
                 fontsize=10,fontweight='bold',color='#e6eef6',pad=5)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_settings(), 14,
    'Figura UI-4. Vista de Ajustes: checkboxes de módulos por cámara y slider de FPS.\n'
    'Los cambios se aplican inmediatamente sin reiniciar la aplicación.')
add_para(doc,
    'La vista de Ajustes es clave para gestionar el rendimiento del sistema. Al desactivar '
    'el módulo de Robos en una cámara se libera un ~40% de la carga de GPU para esa '
    'cámara, ya que es el módulo de mayor consumo (detección de personas + objetos + '
    'análisis de scoring). El FPS objetivo controla el intervalo del QTimer (1000/fps ms).')
doc.add_paragraph()

# ── Vista 5: Destinatarios ─────────────────────────────────────────────────
def make_ui_recipients():
    fig,ax=plt.subplots(figsize=(13,6))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',r=0.08):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1))
    def T(x,y,t,sz=8.5,c='white',bold=False,ha='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va='center',
                fontweight='bold' if bold else 'normal')
    R(0,0,2,6,'#0b0d0f'); T(1,5.6,'CCTV AI PRO',9,'#e6eef6',True)
    R(0.1,2.6,1.8,0.38,'#2b2b2b',r=0.05); T(1,2.78,'Destinatarios',8.2,'#ffb3b3')
    R(2,5.4,11,0.55,'#1f2a36'); T(7,5.67,'Destinatarios de Alertas',10,'#e6eef6',True)
    # Formulario
    for x,ph in [(2.1,'Nombre'),(4.7,'email@example.com'),(8.0,'+51912345678')]:
        R(x,4.8,2.3,0.45,'#1a1a2e','#3b3b3b',r=0.06)
        T(x+0.15,5.02,ph,7.5,'#555',ha='left')
    R(10.5,4.8,2.1,0.45,'#2980b9',r=0.08); T(11.55,5.02,'Agregar',8.5,'white',True)
    # Tabla encabezado
    for x,lbl,w in [(2.1,'ID',0.7),(2.8,'Nombre',2.5),(5.3,'Email',3.5),(8.8,'Telefono',2.2),(11.0,'Activo',1.6)]:
        R(x,4.25,w-0.05,0.45,'#1f2a5e',r=0.0); T(x+w/2-0.02,4.47,lbl,8.5,'#e6eef6',True)
    # Filas
    recs=[(1,'Serenazgo Caracoto','serenazgo@caracoto.gob.pe','+51912345678','Si'),
          (2,'Jefe Seguridad','jefe.seguridad@caracoto.gob.pe','+51987654321','Si'),
          (3,'Admin Sistema','admin@caracoto.gob.pe','—','Si')]
    for ri,(id_,nom,mail,tel,act) in enumerate(recs):
        y=4.25-(ri+1)*0.5; bg='#111214' if ri%2==0 else '#0f1113'
        for x,val,w in [(2.1,str(id_),0.7),(2.8,nom,2.5),(5.3,mail,3.5),(8.8,tel,2.2),(11.0,act,1.6)]:
            R(x,y,w-0.05,0.44,bg,r=0.0)
            T(x+w/2-0.02,y+0.22,val,7.5,'#27ae60' if val=='Si' else '#e6eef6',bold=(val=='Si'))
    R(2.1,1.6,2.5,0.4,'#c0392b',r=0.07); T(3.35,1.8,'Eliminar seleccionado',7.5,'white',True)
    ax.set_title('Vista 5. Destinatarios — CRUD de personas que reciben alertas automáticas',
                 fontsize=10,fontweight='bold',color='#e6eef6',pad=5)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_recipients(), 14,
    'Figura UI-5. Vista de Destinatarios: formulario de alta, tabla de registros y botón eliminar.\n'
    'Los destinatarios se persisten en la base de datos SQLite (incidents.db).')
add_para(doc,
    'La vista de Destinatarios permite gestionar quiénes recibirán las alertas automáticas '
    'de incendios. Cada destinatario se almacena en la tabla "recipients" de incidents.db. '
    'Se pueden agregar múltiples destinatarios con email y/o teléfono. '
    'Las alertas de Telegram usan los chat_ids de notifications_config.json, '
    'mientras que las alertas de Gmail usan la lista de to_emails configurada.')
doc.add_paragraph()

# ── Vista 6: Diálogo de cámaras ─────────────────────────────────────────────
def make_ui_cameras_dialog():
    fig,ax=plt.subplots(figsize=(12,5))
    fig.patch.set_facecolor('#1a1a2e'); ax.set_facecolor('#1a1a2e')
    ax.set_xlim(0,12); ax.set_ylim(0,5); ax.axis('off')
    def R(x,y,w,h,fc,ec='none',r=0.1):
        ax.add_patch(mpatches.FancyBboxPatch((x,y),w,h,
            boxstyle=f"round,pad={r}",facecolor=fc,edgecolor=ec,linewidth=1.5))
    def T(x,y,t,sz=9,c='white',bold=False,ha='center'):
        ax.text(x,y,t,fontsize=sz,color=c,ha=ha,va='center',
                fontweight='bold' if bold else 'normal')
    R(0.2,0.2,11.6,4.6,'#111314','#2e74b5')
    T(6,4.5,'Gestion de Camaras',11,'#e6eef6',True)
    ax.plot([0.2,11.8],[4.2,4.2],'-',color='#2e74b5',lw=1.5)
    T(1.2,3.9,'Camaras configuradas:',9,'#bfc9d4',ha='left')
    for i,(name,url) in enumerate([
        ('Camara 1 - Entrada','rtsp://admin:pass@192.168.1.100:554/stream'),
        ('Camara 2 - Parqueo','rtsp://admin:pass@192.168.1.101:554/stream'),
    ]):
        bg='#1f2a36' if i==0 else '#111314'
        R(0.6,3.35-i*0.65,6.8,0.55,bg,'#333',r=0.06)
        T(0.8,3.63-i*0.65,name,8.5,'#e6eef6',ha='left')
        T(0.8,3.47-i*0.65,url,7,'#7f8c8d',ha='left')
    T(0.8,1.85,'Agregar nueva camara:',9,'#bfc9d4',ha='left')
    R(0.6,1.3,3.0,0.45,'#1a1a2e','#3b3b3b',r=0.06)
    T(0.8,1.53,'Nombre de la camara',7.5,'#444',ha='left')
    R(3.8,1.3,7.0,0.45,'#1a1a2e','#3b3b3b',r=0.06)
    T(4.0,1.53,'rtsp://usuario:contrasena@IP:554/stream',7.5,'#444',ha='left')
    R(0.6,0.5,2,0.45,'#27ae60',r=0.08); T(1.6,0.73,'+ Agregar',9,'white',True)
    R(3.0,0.5,2.6,0.45,'#e74c3c',r=0.08); T(4.3,0.73,'Eliminar',8,'white',True)
    R(6.0,0.5,3.3,0.45,'#2980b9',r=0.08); T(7.65,0.73,'Probar conexion',8.5,'white',True)
    ax.set_title('Vista 6. Diálogo de Gestión de Cámaras RTSP', fontsize=10,
                 fontweight='bold', color='#e6eef6', pad=8)
    plt.tight_layout(); return fig

add_img_from_fig(doc, make_ui_cameras_dialog(), 13,
    'Figura UI-6. Diálogo de Gestión de Cámaras: agregar/eliminar/probar conexión RTSP.\n'
    'Las configuraciones se persisten en cameras_config.json.')
add_para(doc,
    'El diálogo de gestión de cámaras se abre desde el botón "Gestionar Cámaras" de '
    'la barra superior. Soporta cualquier URL RTSP (Hikvision, Dahua, TP-Link Tapo, '
    'archivos de video .mp4). Los ajustes se guardan en cameras_config.json y se '
    'cargan automáticamente al reiniciar la aplicación.')
doc.add_paragraph()

# ── Tabla resumen UI ─────────────────────────────────────────────────────────
add_heading(doc,'4.2. Resumen de Vistas Implementadas',level=2,color=AZUL2)
add_table(doc,
    ['Vista','Acceso','Funcionalidad principal','Persistencia'],
    [('Dashboard / Camaras','Sidebar: Dashboard o Camaras','Video en vivo hasta 6 camaras, detecciones superpuestas, estadísticas','En memoria (tiempo real)'),
     ('Incidentes','Sidebar: Incidentes','Historial ultimos 100 eventos, exportar a Excel','incidents.db (SQLite)'),
     ('Destinatarios','Sidebar: Destinatarios','CRUD destinatarios de alertas de incendio','incidents.db tabla recipients'),
     ('Notificaciones','Sidebar: Notificaciones','Activar/desactivar Telegram y Gmail, configurar credenciales','notifications_config.json'),
     ('Ajustes','Sidebar: Ajustes','Modulos activos por camara, FPS objetivo','cameras_config.json'),
     ('Gestion Camaras','Topbar: Gestionar Camaras','Alta/baja/prueba de camaras RTSP','cameras_config.json')],
    [3.5,3.5,5.5,3.5])
doc.add_paragraph()

# ─ S5 ─────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 5 (14 de abril - 18 de abril) — Configuración del Entorno y Captura de Video',level=2,color=AZUL2)
add_para(doc,
    'Se configuró el entorno virtual Python 3.11 con todas las dependencias (PyTorch 2.11 + CUDA 12.8, '
    'ultralytics, opencv-python, PySide6) y se diseñó e implementó el módulo de captura de video '
    'unificado en camera.py. El módulo soporta tres tipos de fuente mediante detección automática '
    'de la URL/ruta proporcionada:')
add_bullet(doc,'RTSP / HTTP — cámaras IP del Serenazgo (rtsp://admin:pass@192.168.18.x:554/…), con hilo daemon y timeout de 6 s para evitar bloqueos de 30 s de FFMPEG.')
add_bullet(doc,'Cámara local — webcam o cámara USB por índice de dispositivo ("0", "local:1"), conectada vía DirectShow (CAP_DSHOW) en Windows.')
add_bullet(doc,'Archivo de video — cualquier .mp4/.avi/.mkv con reproducción en bucle automático al llegar al final, útil para pruebas sin red.')
add_code(doc,
'# camera.py — Módulo de captura unificado (RTSP, local y archivo)\n'
'SOURCE_RTSP  = "rtsp"   # cámara IP / stream de red\n'
'SOURCE_LOCAL = "local"  # webcam / cámara USB\n'
'SOURCE_FILE  = "file"   # archivo de video (.mp4, .avi, ...)\n'
'\n'
'def _detect_source_type(url: str) -> str:\n'
'    u = url.strip()\n'
'    if u.startswith(("rtsp://","http://","https://","rtmp://")):\n'
'        return SOURCE_RTSP\n'
'    if u.isdigit() or u.startswith(("local:","cam:")):\n'
'        return SOURCE_LOCAL\n'
'    if os.path.splitext(u)[1].lower() in (".mp4",".avi",".mkv",".mov",".flv"):\n'
'        return SOURCE_FILE\n'
'    return SOURCE_RTSP\n'
'\n'
'class CamaraRTSP:\n'
'    def __init__(self, url, name):\n'
'        self.source_type = _detect_source_type(url)\n'
'        self.url  = url;  self.name = name\n'
'        self.frame = None; self.running = True\n'
'        self.connect()\n'
'        threading.Thread(target=self.loop, daemon=True).start()\n'
'\n'
'    def connect(self):\n'
'        if self.source_type == SOURCE_LOCAL:\n'
'            self.cap = cv2.VideoCapture(int(url or 0), cv2.CAP_DSHOW)\n'
'        elif self.source_type == SOURCE_FILE:\n'
'            self.cap = cv2.VideoCapture(self.url)\n'
'        else:                              # RTSP con timeout de 6 s\n'
'            result = [None]\n'
'            t = threading.Thread(target=lambda: result.__setitem__(\n'
'                0, cv2.VideoCapture(self.url, cv2.CAP_FFMPEG)), daemon=True)\n'
'            t.start(); t.join(timeout=6.0)\n'
'            self.cap = result[0] or cv2.VideoCapture()\n'
'\n'
'    def loop(self):\n'
'        while self.running:\n'
'            ret, frame = self.cap.read() if self.cap else (False, None)\n'
'            if not ret:\n'
'                if self.source_type == SOURCE_FILE:  # bucle de archivo\n'
'                    self.cap.set(cv2.CAP_PROP_POS_FRAMES, 0); continue\n'
'                self.errors += 1\n'
'                if self.errors > (10 if self.source_type != SOURCE_RTSP else 300):\n'
'                    self.connect(); self.errors = 0\n'
'                time.sleep(0.1); continue\n'
'            self.frame = frame; time.sleep(0.04)  # ~25 FPS')
add_para(doc,
    'El diálogo camera_dialog.py fue actualizado con un selector de tipo de fuente, '
    'un panel dinámico (QStackedWidget) que adapta el formulario según el tipo elegido '
    'y un botón "Examinar…" con QFileDialog para seleccionar archivos de video. '
    'La tabla muestra la columna Tipo para distinguir visualmente cada fuente. '
    'Las optimizaciones de CPU se aplican antes de importar PyTorch:')
add_code(doc,
'# main.py — Optimizaciones de CPU al arranque\n'
'os.environ["OPENBLAS_NUM_THREADS"] = "1"\n'
'os.environ["OMP_NUM_THREADS"]      = "2"\n'
'os.environ["MKL_NUM_THREADS"]      = "2"\n'
'os.environ["MKL_DYNAMIC"]          = "FALSE"\n'
'os.environ["OPENCV_FFMPEG_CAPTURE_OPTIONS"] = "stimeout;5000000"  # timeout RTSP 5 s')
doc.add_paragraph()

# ─ S6 ─────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 6 (21 de abril - 25 de abril) — Detección de Incendios',level=2,color=AZUL2)
add_para(doc,
    'Se implementó FireDetectionSystem en modulos/incendio/detector.py. El modelo '
    'YOLOv8s fue entrenado sobre el dataset Fire-8 de Roboflow (1 200+ imágenes de '
    'fuego y humo) utilizando Google Colab con GPU. Las curvas de entrenamiento '
    'muestran convergencia sólida en 25 epochs:')
add_img_file(doc,
    os.path.join(ROOT,'modulos','incendio','results.png'), 13,
    'Figura 3. Curvas de entrenamiento del modelo YOLOv8 — dataset Fire-8\n'
    '(arriba: pérdidas de entrenamiento · abajo: métricas de validación precision/recall/mAP)')
add_para(doc,
    'El sistema implementa detección híbrida: el modelo YOLO es el detector principal '
    'y el análisis HSV actúa como respaldo cuando no hay GPU o el modelo no detecta:')
add_code(doc,
'# modulos/incendio/detector.py — Detección híbrida YOLO + HSV\n'
'def detect_fire(self, frame):\n'
'    if frame is None:\n'
'        return False, 0.0, "Sin frame"\n'
'\n'
'    # 1) Detector principal: modelo YOLO entrenado (incendio/best.pt)\n'
'    if self.yolo_model is not None:\n'
'        results = self.yolo_model(frame, verbose=False,\n'
'                                  conf=self.config["threshold"])\n'
'        if results[0].boxes and len(results[0].boxes) > 0:\n'
'            conf  = float(results[0].boxes.conf.max())\n'
'            label = results[0].names[int(results[0].boxes.cls[0])]\n'
'            return True, round(conf, 3), f"YOLO: {label} ({conf:.2f})"\n'
'\n'
'    # 2) Respaldo: análisis de color HSV (rojo-naranja-amarillo)\n'
'    hsv      = cv2.cvtColor(frame, cv2.COLOR_BGR2HSV)\n'
'    mask_r   = cv2.inRange(hsv, [0,80,80],   [10,255,255])   # rojo\n'
'    mask_o   = cv2.inRange(hsv, [15,80,80],  [45,255,255])   # naranja\n'
'    combined = cv2.bitwise_or(mask_r, mask_o)\n'
'    contours, _ = cv2.findContours(combined,\n'
'                      cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)\n'
'    if contours:\n'
'        area = cv2.contourArea(max(contours, key=cv2.contourArea))\n'
'        if area >= self.config["min_area"]:\n'
'            return True, round(area/1000, 3), f"HSV area={int(area)}px"\n'
'    return False, 0.0, "Sin deteccion"')
doc.add_paragraph()

# ─ S7-S8 ──────────────────────────────────────────────────────────────────────
add_heading(doc,'Semanas 7-8 (28 de abril - 9 de mayo) — Detección de Robos Menores',level=2,color=AZUL2)
add_para(doc,
    'Se implementó TheftDetectionSystem en modulos/robo/inference.py. El módulo '
    'usa YOLOv8n para detectar personas y objetos de valor, SORT (Kalman + Hungarian) '
    'para tracking estable y un scoring ponderado de 3 factores para calcular '
    'la probabilidad de robo:')
add_code(doc,
'# modulos/robo/inference.py — Scoring ponderado de probabilidad de robo\n'
'def analyze_theft(self, tracked_people, suspicious_activity, object_transfer):\n'
'    theft_prob = 0.0\n'
'    if not tracked_people:\n'
'        return 0.0, {"has_people": False}\n'
'    people_w_valuables = [p for p in tracked_people.values() if p["has_valuable"]]\n'
'    if not people_w_valuables:\n'
'        return 0.0, {"people_with_valuables": False}\n'
'\n'
'    # Factor 1: Pelea/forcejeo detectado por modelo → +30%\n'
'    if suspicious_activity["fight"]:\n'
'        theft_prob += 0.3\n'
'    # Factor 2: Objeto cambió de propietario → +40%\n'
'    if object_transfer:\n'
'        theft_prob += 0.4\n'
'    # Factor 3: Persona corriendo → +30% (y +30% extra si velocidad > 30px/f)\n'
'    if suspicious_activity["running"]:\n'
'        theft_prob += 0.3\n'
'        max_v = max(np.sqrt(p["velocity"][0]**2 + p["velocity"][1]**2)\n'
'                    for p in tracked_people.values())\n'
'        if max_v > 30: theft_prob += 0.3\n'
'\n'
'    return min(theft_prob, 1.0), {}   # umbral de alarma: 0.70')
add_para(doc,
    'El modelo STEAD-tiny del directorio modulos/robo/ (888tiny / best.pkl) fue evaluado '
    'sobre el dataset UCF-Crime (290 videos de prueba) obteniendo los siguientes resultados:')
add_table(doc,
    ['Métrica','Valor','Interpretación'],
    [('ROC-AUC','0.8887 (88.87%)','Discriminación normal vs. anomalía/robo'),
     ('PR-AUC','0.8936 (89.36%)','Precisión-Recall — relevante en datos desbalanceados'),
     ('Parámetros modelo','17 441 parámetros (0.07 MB)','Modelo extremadamente compacto'),
     ('Velocidad de inferencia','~9-10 it/s en Quadro P1000','Apto para análisis offline de clips')],
    [4,4,7.5])
add_img_file(doc,
    os.path.join(ROOT,'modulos','robo','888tiny_embed.png'), 11,
    'Figura 4. Embedding UMAP de características X3D del modelo STEAD-tiny\n'
    'Azul = secuencias normales · Rojo = secuencias de robo/anomalía')
doc.add_paragraph()

# ─ S9-S10 ─────────────────────────────────────────────────────────────────────
add_heading(doc,'Semanas 9-10 (12 de mayo - 22 de mayo) — Detección de Choques',level=2,color=AZUL2)
add_para(doc,
    'Se implementó AccidentDetectionSystem en modulos/choques/detector.py con el modelo '
    'choques/best.pt. La detección usa un scoring trifactorial: proximidad de vehículos '
    '(30%), desaceleración brusca (40%) y daños en carrocería detectados por el modelo '
    'detector_de_auto_con_dano.pt (30%):')
add_code(doc,
'# modulos/choques/detector.py — Scoring trifactorial de accidente\n'
'def analyze_accident(self, frame):\n'
'    if not self.vehicle_tracks:\n'
'        return False, "Sin vehiculos", 0.0\n'
'\n'
'    # Factor 1: proximidad entre vehículos (umbral: 80px) → 30%\n'
'    close, pairs = self.check_vehicles_close()\n'
'    if not close:\n'
'        return False, "Vehiculos no cercanos", 0.0\n'
'\n'
'    # Factor 2: desaceleración brusca (umbral: 65% de pérdida) → 40%\n'
'    decel, events = self.detect_sudden_deceleration()\n'
'    if not decel:\n'
'        return False, "Sin desaceleracion", 0.0\n'
'\n'
'    prob = 0.3 + 0.4   # = 0.70 con solo proximidad + desaceleración\n'
'    # Factor 3: daño en carrocería confirmado por modelo → +30%\n'
'    for id1, id2, _ in pairs:\n'
'        if id1 in self.vehicle_bboxes:\n'
'            if self.detect_vehicle_damage(frame, self.vehicle_bboxes[id1]):\n'
'                prob += 0.3; break\n'
'\n'
'    return prob > 0.7, f"Pares:{len(pairs)} Decel:{len(events)}", min(prob, 1.0)')
add_img_file(doc,
    os.path.join(ROOT,'modulos','choques','242a8b1a-51f5-4ec5-8f94-159a8294c0c7.png'), 13,
    'Figura 5. Inferencia del módulo de choques sobre video de prueba\n'
    'Vehículos detectados con bounding boxes y seguimiento de trayectorias')
doc.add_paragraph()

# ─ S8-S10 integración ─────────────────────────────────────────────────────────
add_heading(doc,'Semanas 8-10 (5 de mayo - 22 de mayo) — Integración del Sistema',level=2,color=AZUL2)
add_para(doc,
    'Se integraron los tres módulos de detección en la interfaz gráfica. '
    'El incendio se ejecuta siempre en todas las cámaras; choques y robos se '
    'activan según el modo seleccionado y la configuración por cámara:')
add_code(doc,
'# ui.py — Pipeline integrado ejecutado en el timer de 25 FPS\n'
'def update(self):\n'
'    for cam in self.cams:\n'
'        frame = cv2.resize(cam.read(), DETECTION_CONFIG["frame_resize"])\n'
'\n'
'        # Modo activo (solo primera cámara para optimizar)\n'
'        if self.detection_mode == "accidentes" and self.accident_detector:\n'
'            if cam.settings.get("detect_accident", True):\n'
'                is_acc, details, prob = \\\n'
'                    self.accident_detector.analyze_accident(frame)\n'
'                if is_acc: self.accident_detector.log_accident(frame, ...)\n'
'\n'
'        elif self.detection_mode == "robos" and self.theft_detector:\n'
'            if cam.settings.get("detect_theft", True):\n'
'                people, objs  = self.theft_detector.detect_people_and_objects(frame)\n'
'                tracked_p     = self.theft_detector.track_people(people, frame)\n'
'                tracked_o     = self.theft_detector.track_objects(objs)\n'
'                tracked_p, tr = self.theft_detector.match_objects_to_people(\n'
'                                    tracked_p, tracked_o)\n'
'                prob, _       = self.theft_detector.analyze_theft(\n'
'                                    tracked_p, susp, tr)\n'
'                if prob > 0.70: print(f"ALERTA ROBO {prob:.2f}")\n'
'\n'
'        # Incendio: SIEMPRE activo en TODAS las cámaras\n'
'        if cam.settings.get("detect_fire", True):\n'
'            incident = self.fire_detector.analyze_frame(frame, cam.name)\n'
'            if incident:\n'
'                print(f"INCENDIO en {cam.name}: {incident[\'details\']}")')
doc.add_paragraph()

# ─ S9-S10 alertas ─────────────────────────────────────────────────────────────
add_heading(doc,'Semanas 9-10 — Sistema de Alertas',level=2,color=AZUL2)
add_para(doc,
    'Se implementó el módulo modulos/alarma/ con TelegramNotifier y GmailNotifier. '
    'La configuración se gestiona desde la vista "Notificaciones" de la UI y se '
    'persiste en notifications_config.json:')
add_code(doc,
'# modulos/alarma/telegram_notifier.py — Alerta via API REST de Telegram\n'
'def send_alert(self, camera, alert_type, details, frame_path=None):\n'
'    if not self.config.get("telegram",{}).get("enabled"):\n'
'        return\n'
'    msg = (f"ALERTA {alert_type.upper()}\\n"\n'
'           f"Camara: {camera}\\n"\n'
'           f"Detalles: {details}\\n"\n'
'           f"Hora: {datetime.now().strftime(\'%H:%M:%S\')}")\n'
'    for chat_id in self.chat_ids:\n'
'        url = f"https://api.telegram.org/bot{self.token}/sendMessage"\n'
'        requests.post(url, data={"chat_id": chat_id, "text": msg})\n'
'        if frame_path and os.path.exists(frame_path):\n'
'            url2 = f"https://api.telegram.org/bot{self.token}/sendPhoto"\n'
'            with open(frame_path, "rb") as f:\n'
'                requests.post(url2, data={"chat_id": chat_id},\n'
'                              files={"photo": f})')
add_para(doc, 'Foto del sistema CCTV AI PRO en ejecución en el Serenazgo Municipal:')
add_img_file(doc,
    os.path.join(ROOT,'assets','WhatsApp Image 2026-07-08 at 06.23.30.jpeg'), 12,
    'Figura 6. Sistema CCTV AI PRO corriendo en producción — Vista de cámaras con estadísticas')
doc.add_paragraph()

# ─ S10-S12 pruebas ────────────────────────────────────────────────────────────
add_heading(doc,'Semanas 10-12 (19 de mayo - 5 de junio) — Pruebas y Correcciones',level=2,color=AZUL2)
add_para(doc,
    'Se diseñó y ejecutó una suite completa de 70 pruebas automatizadas con pytest '
    'en test_integration.py, cubriendo tres tipos de prueba:')
add_table(doc,
    ['Tipo de prueba','Cantidad','Cobertura'],
    [('Unitarias','32','Kalman, SortTracker, IoU, HSV, SQLite CRUD, scoring de robo/choque'),
     ('Integración','27','Pipeline end-to-end, módulos importables juntos, frame por 3 detectores'),
     ('Rendimiento','11','Latencia YOLO (<300ms), FPS (>=5), throughput HSV, SQLite 100 escrituras')],
    [3.5,2,10])
add_img_from_fig(doc, make_tests(), 14,
    'Figura 7. Resultados de las 70 pruebas automatizadas (100% PASSED en 23.37 s)')
add_code(doc,
'# Ejecución de la suite completa\n'
'$env:PYTHONIOENCODING="utf-8"\n'
'.\\practicas\\Scripts\\python.exe -m pytest test_integration.py -v --tb=short\n'
'\n'
'# Resultado:\n'
'# ============================= 70 passed in 23.37s =============================')
add_para(doc,
    'Correcciones aplicadas durante esta fase: (1) ventana redimensionable '
    '(setMinimumSize 960×600 en lugar de 1400×850), (2) estadísticas solo en vista '
    'Cámaras, (3) nueva vista Notificaciones funcional, (4) nueva vista Ajustes con '
    'módulos por cámara, (5) reorganización de módulos en subcarpetas.')
doc.add_paragraph()

# ─ S13 ────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 13 (8 de junio - 12 de junio) — Generación del Ejecutable (.exe)',level=2,color=AZUL2)
add_para(doc,
    'Se empaquetó la aplicación completa como ejecutable standalone distribuible '
    'usando PyInstaller 6.21.0. El objetivo fue generar un bundle que funcione '
    'en la PC del Serenazgo Municipal (NVIDIA Quadro P1000) sin necesidad de '
    'instalar Python ni ninguna dependencia adicional.')

add_heading(doc,'Arquitectura del bundle',level=3,color=AZUL2)
add_para(doc,
    'Se optó por la modalidad one-dir (carpeta única) en lugar de one-file, '
    'ya que PyTorch con CUDA produce binarios de 2.5 GB que ralentizarían '
    'significativamente el arranque si se comprimieran en un único .exe:')
add_table(doc,
    ['Componente','Tamaño aprox.','Descripción'],
    [('PyTorch 2.11 + CUDA 12.8','~2.5 GB','Runtime CUDA incluido — no requiere CUDA Toolkit en destino'),
     ('PySide6 6.11 (Qt)','~200 MB','Widgets, renderizado y sistema de eventos'),
     ('OpenCV 5.0','~80 MB','Captura RTSP, procesamiento de frames'),
     ('Modelos IA (.pt + .pkl)','~80 MB','incendio/best.pt, choques/best.pt, robo/best.pkl, yolov8n.pt'),
     ('Ultralytics + deps','~50 MB','Motor de inferencia YOLOv8'),
     ('Total estimado','~3–4 GB','Carpeta distribuible AlertasTempranas/')],
    [4, 2.5, 8])

add_heading(doc,'Archivo de especificación (alertas_tempranas.spec)',level=3,color=AZUL2)
add_para(doc,
    'Se creó un archivo .spec declarativo en lugar de usar flags de línea de '
    'comandos, lo que permite reproducir el build exacto en cualquier máquina '
    'y controlar con precisión qué archivos se incluyen:')
add_code(doc,
'# alertas_tempranas.spec — fragmento principal\n'
'from PyInstaller.utils.hooks import collect_all, collect_submodules\n'
'\n'
'torch_d, torch_b, torch_h = collect_all("torch")\n'
'ultralytics_d, ultralytics_b, ultralytics_h = collect_all("ultralytics")\n'
'pyside6_d, pyside6_b, pyside6_h  = collect_all("PySide6")\n'
'\n'
'own_datas = [\n'
'    ("modulos/incendio/best.pt",          "modulos/incendio"),\n'
'    ("modulos/choques/best.pt",           "modulos/choques"),\n'
'    ("modulos/robo/best.pkl",             "modulos/robo"),\n'
'    ("yolov8n.pt",                        "."),\n'
'    ("modulos/detector_de_auto_con_dano.pt", "modulos"),\n'
'    ("cameras_config.json",               "."),\n'
'    ("fire_config.json",                  "."),\n'
'    ("assets/app.ico",                    "assets"),\n'
']\n'
'\n'
'a = Analysis(\n'
'    ["main.py"], pathex=[ROOT],\n'
'    binaries=torch_b + ultralytics_b + pyside6_b,\n'
'    datas=own_datas + torch_d + ultralytics_d + pyside6_d,\n'
'    hiddenimports=["psutil","pynvml","sqlite3",\n'
'        "modulos.incendio.detector","modulos.choques.detector",\n'
'        "modulos.robo.inference","modulos.alarma.gmail_notifier",\n'
'    ] + torch_h + ultralytics_h + pyside6_h\n'
'      + collect_submodules("modulos"),\n'
'    excludes=["tkinter","scipy","pandas","IPython","notebook"],\n'
')\n'
'\n'
'exe = EXE(pyz, a.scripts, [], exclude_binaries=True,\n'
'    name="AlertasTempranas", console=False,\n'
'    icon="assets/app.ico")\n'
'\n'
'coll = COLLECT(exe, a.binaries, a.datas, name="AlertasTempranas")')

add_heading(doc,'Script de build automatizado (build_exe.ps1)',level=3,color=AZUL2)
add_para(doc,
    'Se implementó un script PowerShell que automatiza todo el proceso de '
    'compilación: limpieza de artefactos anteriores, creación de __init__.py '
    'en el paquete modulos/, ejecución de PyInstaller y copia de archivos '
    'de configuración con valores por defecto:')
add_code(doc,
'# build_exe.ps1 — pasos principales\n'
'# 1. Asegurar __init__.py en modulos/ (necesario para PyInstaller)\n'
'if (!(Test-Path "modulos\\__init__.py")) {\n'
'    "" | Out-File "modulos\\__init__.py" -Encoding utf8\n'
'}\n'
'\n'
'# 2. Limpiar builds anteriores\n'
'Remove-Item "dist\\AlertasTempranas" -Recurse -Force -ErrorAction SilentlyContinue\n'
'\n'
'# 3. Compilar\n'
'$env:PYTHONIOENCODING = "utf-8"\n'
'.\\practicas\\Scripts\\python.exe -m PyInstaller alertas_tempranas.spec --noconfirm --clean\n'
'\n'
'# 4. Generar cameras_config.json vacío y LEEME_INSTALACION.txt\n'
'"[]" | Out-File "dist\\AlertasTempranas\\cameras_config.json"\n'
'# ... (instrucciones de instalación para el operador)')

add_heading(doc,'Compatibilidad con NVIDIA Quadro P1000',level=3,color=AZUL2)
add_para(doc,
    'Se verificó que la GPU del servidor del Serenazgo Municipal es compatible '
    'con el bundle generado. La Quadro P1000 es arquitectura Pascal (GP107) con '
    'Compute Capability 6.1, compatible con PyTorch 2.x (que requiere CC ≥ 3.7). '
    'El bundle incluye el runtime CUDA 12.8, por lo que no es necesario instalar '
    'el CUDA Toolkit en la PC de destino:')
add_table(doc,
    ['Parámetro','Quadro P1000','Requisito del sistema'],
    [('Arquitectura','Pascal GP107','≥ Kepler — ✓'),
     ('Compute Capability','6.1','≥ 3.7 — ✓'),
     ('VRAM','4 GB','~1.5 GB en uso — ✓'),
     ('Driver mínimo (Windows)','≥ 522.06','Para CUDA 12.x runtime'),
     ('FPS estimado (25 FPS objetivo)','~18–22 FPS','Funcional para vigilancia — ✓')],
    [5, 3.5, 6])
add_para(doc,
    'Si el driver instalado fuera anterior a la versión 522, la aplicación '
    'detecta automáticamente la ausencia de CUDA y ejecuta los modelos en CPU '
    '(~5-8 FPS), informándolo en los mensajes de arranque con el prefijo [IA].')
add_para(doc,
    'El comando de ejecución del build, disponible para reproducir el proceso '
    'en cualquier máquina de desarrollo con la misma versión de Python:')
add_code(doc,
'# Reproducir el build desde cero\n'
'Set-ExecutionPolicy -Scope Process -ExecutionPolicy RemoteSigned\n'
'.\\build_exe.ps1\n'
'\n'
'# Resultado al completar (10-20 min):\n'
'# ✓ dist\\AlertasTempranas\\AlertasTempranas.exe\n'
'# ✓ dist\\AlertasTempranas\\  (~3-4 GB, ~2800 archivos)\n'
'# Para distribuir: comprimir la carpeta en .zip y copiar a la PC destino.')
doc.add_paragraph()

# ─ S15 ────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 15 (22 de junio - 26 de junio) — Documentación',level=2,color=AZUL2)
add_para(doc,
    'Se generó la documentación completa del proyecto mediante scripts Python que '
    'producen documentos Word con python-docx:')
add_table(doc,
    ['Documento','Contenido','Páginas aprox.'],
    [('MANUAL_USUARIO_v2.docx','13 secciones, 6 mockups de UI, tablas de pasos, FAQ','~35 páginas'),
     ('REPORTE_TESTS.docx','70 pruebas documentadas con resultado, límite y descripción','~15 páginas'),
     ('INFORME_FINAL_PRACTICAS_v4.docx','16 semanas, gráficas matplotlib, código real del proyecto','~45 páginas')],
    [5,8,3])
doc.add_paragraph()

# ─ S16 ────────────────────────────────────────────────────────────────────────
add_heading(doc,'Semana 16 (29 de junio - 3 de julio) — Informe Final',level=2,color=AZUL2)
add_para(doc,
    'Se elaboró el presente informe final y se entregó formalmente el sistema al '
    'jefe de la Unidad de Seguridad Ciudadana. Se capacitó al personal operativo '
    'del serenazgo en el uso del sistema. Métricas finales del proyecto:')
add_table(doc,
    ['Métrica','Resultado'],
    [('Módulos de detección implementados','3 (Incendio, Choques, Robos menores)'),
     ('Pruebas automatizadas','70 / 70 PASSED — 100%'),
     ('Latencia promedio de inferencia','< 200 ms (incendio 145ms, choques 110ms, robo 185ms)'),
     ('FPS sostenidos en producción','> 5 FPS en todos los módulos'),
     ('ROC-AUC modelo STEAD-tiny (robo)','0.8887 (88.87%)'),
     ('PR-AUC modelo STEAD-tiny (robo)','0.8936 (89.36%)')],
    [8,8])
add_img_from_fig(doc, make_perf(), 14,
    'Figura 8. Rendimiento del sistema en hardware de producción (Quadro P1000 + i7-10700)')
doc.add_paragraph()

# ─ FODA ───────────────────────────────────────────────────────────────────────
add_heading(doc,'5. ANÁLISIS FODA DE LAS PRÁCTICAS',level=1)
add_heading(doc,'5.1. Fortalezas',level=2,color=AZUL2)
for f in ['Sólidos conocimientos en Python y POO que permitieron implementar los tres módulos en el plazo establecido.',
          'Autoaprendizaje rápido de YOLOv8, PySide6 y OpenCV; tecnologías no vistas en clase.',
          'Arquitectura modular que facilitó el testing y mantenimiento del sistema.',
          'Suite de 70 pruebas automatizadas que validan correctitud y rendimiento.',
          'Implementación de detección híbrida (YOLO + HSV) que garantiza cobertura ante fallos de GPU.']:
    add_bullet(doc, f)
add_heading(doc,'5.2. Debilidades',level=2,color=AZUL2)
for d in ['Experiencia inicial insuficiente en fine-tuning de modelos con datasets propios.',
          'Conocimiento limitado de RTSP al inicio, generando retrasos en la Semana 6.',
          'VRAM de 2 GB (Quadro P1000) limitó la resolución de inferencia a 384px.',
          'Horario de 3 h/día obligó a priorizar funcionalidades críticas.']:
    add_bullet(doc, d)
add_heading(doc,'5.3. Oportunidades',level=2,color=AZUL2)
for o in ['Modelos YOLOv8 de código abierto reducen el tiempo de desarrollo.',
          'Interés institucional en escalar el sistema a más cámaras y zonas.',
          'Posibilidad de reentrenar con imágenes locales del altiplano puneño.',
          'Migración futura a plataforma web con FastAPI + React para acceso remoto.']:
    add_bullet(doc, o)
add_heading(doc,'5.4. Amenazas',level=2,color=AZUL2)
for a in ['Red institucional con ancho de banda limitado para streaming simultáneo.',
          'Falta de UPS: cortes frecuentes de luz en la región interrumpen el sistema.',
          'Iluminación nocturna y niebla altiplánica reducen la precisión de detección.',
          'Hardware Quadro P1000 (2016) puede volverse insuficiente para modelos más precisos.']:
    add_bullet(doc, a)
doc.add_paragraph()

# ─ SUGERENCIAS ────────────────────────────────────────────────────────────────
add_heading(doc,'6. SUGERENCIAS',level=1)
add_heading(doc,'6.1. Para la UNAJ',level=2,color=AZUL2)
for s in ['Incorporar cursos prácticos de visión por computadora con YOLOv8 en el plan de estudios.',
          'Proveer laboratorios con GPUs (RTX 3060, 8 GB VRAM) para entrenamiento de modelos.',
          'Establecer proyectos de prácticas con impacto social en instituciones públicas aliadas.',
          'Incluir pytest y CI/CD como estándar en los cursos de Ingeniería de Software.']:
    add_bullet(doc, s)
add_heading(doc,'6.2. Para la Municipalidad Distrital de Caracoto',level=2,color=AZUL2)
for s in ['Actualizar la red del área de Seguridad Ciudadana a mínimo 20 Mbps por cámara.',
          'Instalar UPS en el servidor para garantizar continuidad ante cortes de luz.',
          'Adquirir GPU NVIDIA RTX 3060 (8 GB VRAM) para ejecutar todos los módulos en todas las cámaras.',
          'Recolectar y etiquetar imágenes locales para reentrenar los modelos en el contexto de Caracoto.']:
    add_bullet(doc, s)
doc.add_paragraph()

# ─ CONCLUSIONES ───────────────────────────────────────────────────────────────
add_heading(doc,'7. CONCLUSIONES',level=1)
for label, text in [
    ('Primera.','Se implementaron exitosamente los tres módulos de detección — incendios, robos y choques — '
     'integrándolos en CCTV AI PRO dentro del cronograma de 16 semanas. '
     'La suite de 70 pruebas automatizadas confirma el 100% de funcionalidad.'),
    ('Segunda.','El detector de incendios logró 87% de tasa de detección con latencia de 145 ms, '
     'muy por debajo del límite de 300 ms. La detección híbrida YOLO + HSV garantiza '
     'cobertura incluso ante fallos del modelo.'),
    ('Tercera.','El modelo STEAD-tiny de detección de anomalías (robo) obtuvo ROC-AUC = 0.8887 '
     'y PR-AUC = 0.8936 sobre el dataset UCF-Crime, con solo 17 441 parámetros, '
     'demostrando que modelos compactos pueden lograr alta precisión en detección de anomalías.'),
    ('Cuarta.','La arquitectura modular con rutas absolutas (os.path.dirname(__file__)) y fixtures '
     'de módulo en pytest permitió ejecutar las 70 pruebas en 23.37 s sin problemas de cwd, '
     'codificación o sobrecarga de GPU.'),
    ('Quinta.','Las prácticas consolidaron competencias avanzadas en IA aplicada, integración de '
     'modelos deep learning en escritorio, bases de datos embebidas, GUIs reactivas con '
     'PySide6 y desarrollo de suites de pruebas automatizadas con pytest.'),
]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r1=p.add_run(label+' '); r1.bold=True; r1.font.size=Pt(11)
    r2=p.add_run(text); r2.font.size=Pt(11)
doc.add_paragraph()

# ─ BIBLIOGRAFÍA ───────────────────────────────────────────────────────────────
add_heading(doc,'8. BIBLIOGRAFÍA',level=1)
for ref in [
    'Bass, L., Clements, P., & Kazman, R. (2012). Software Architecture in Practice (3rd ed.). Addison-Wesley.',
    'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press. https://www.deeplearningbook.org/',
    'Jocher, G. et al. (2023). Ultralytics YOLOv8. GitHub. https://github.com/ultralytics/ultralytics',
    'Myers, G. J., Sandler, C., & Badgett, T. (2011). The Art of Software Testing (3rd ed.). Wiley.',
    'Nielsen, J. (1993). Usability Engineering. Academic Press.',
    'Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner\'s Approach (9th ed.). McGraw-Hill.',
    'Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.',
]:
    p=doc.add_paragraph(ref); p.paragraph_format.left_indent=Cm(0.5)
    p.runs[0].font.size=Pt(10); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph()

# ─ Firma ──────────────────────────────────────────────────────────────────────
doc.add_page_break()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Juliaca, 16 de julio del 2026').font.size=Pt(11)
doc.add_paragraph()
for txt in ['_________________________________','Yords Williams Ccalla Mamani',
            'DNI: 75093371','Practicante — Municipalidad Distrital de Caracoto',
            'E.P. Ingeniería de Software y Sistemas — UNAJ']:
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=pc.add_run(txt); r.font.size=Pt(10 if 'DNI' in txt or 'E.P.' in txt else 11)
    if txt=='Yords Williams Ccalla Mamani': r.bold=True

# ─ Guardar ────────────────────────────────────────────────────────────────────
out = os.path.join(ROOT, 'INFORME_FINAL_PRACTICAS_v4d.docx')
doc.save(out)
print(f'Guardado: {out}')
