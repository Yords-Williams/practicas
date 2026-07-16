"""Genera MANUAL_USUARIO.docx para CCTV AI PRO."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3);  s.right_margin = Cm(2.5)

AZUL  = RGBColor(0x1F, 0x2A, 0x5E)
AZUL2 = RGBColor(0x2E, 0x74, 0xB5)
GRIS  = RGBColor(0x44, 0x44, 0x44)
VERDE = RGBColor(0x1E, 0x8B, 0x4C)

def shd(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)

def h(text, level=1, color=AZUL):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.color.rgb = color; r.bold = True
    return p

def p(text, size=11, bold=False, italic=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; return para

def c(text, bold=False, size=12, color=None):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return para

def note(text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run("📌 " + text)
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = AZUL2
    return para

def step_table(steps):
    """Tabla de pasos numerados."""
    t = doc.add_table(rows=len(steps)+1, cols=2); t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Paso"; t.rows[0].cells[1].text = "Acción"
    for c_,r_ in [(t.rows[0].cells[0], 'paso'), (t.rows[0].cells[1], 'paso')]:
        c_.paragraphs[0].runs[0].bold = True
        c_.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd(c_, '2E74B5')
    for i, (num, text) in enumerate(steps):
        t.rows[i+1].cells[0].text = str(num)
        t.rows[i+1].cells[1].text = text
        t.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i+1].cells[0].paragraphs[0].runs[0].font.color.rgb = AZUL2
        t.rows[i+1].cells[0].width = Cm(1.5); t.rows[i+1].cells[1].width = Cm(14)
        if i % 2 == 0:
            shd(t.rows[i+1].cells[0], 'F0F5FB'); shd(t.rows[i+1].cells[1], 'F0F5FB')

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
c('UNIVERSIDAD NACIONAL DE JULIACA', bold=True, size=15, color=AZUL)
c('Facultad de Ciencias de Ingenierías', bold=True, size=12, color=AZUL)
c('Escuela Profesional de Ingeniería de Software y Sistemas', size=11, color=AZUL)
doc.add_paragraph()
c('MANUAL DE USUARIO', bold=True, size=16, color=AZUL)
c('CCTV AI PRO — Sistema de Monitoreo Inteligente', bold=True, size=13, color=AZUL2)
c('Municipalidad Distrital de Caracoto', size=11, color=GRIS)
doc.add_paragraph()
t = doc.add_table(rows=4, cols=2); t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([('Versión:','1.0'),('Fecha:','Julio 2026'),
                            ('Elaborado por:','Yords Williams Ccalla Mamani'),
                            ('Área:','Seguridad Ciudadana — Serenazgo Municipal')]):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    for j in range(2): t.rows[i].cells[j].width = Cm(7)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════════
h('ÍNDICE DE CONTENIDO', level=1)
indice = [
    "1. Introducción y descripción del sistema",
    "2. Requisitos mínimos del sistema",
    "3. Inicio de la aplicación",
    "4. Interfaz principal — Vista de Cámaras",
    "5. Menú lateral — Navegación",
    "6. Vista de Incidentes",
    "7. Vista de Destinatarios",
    "8. Vista de Notificaciones — activar Telegram y Gmail",
    "9. Vista de Ajustes — módulos por cámara y FPS",
    "10. Gestión de cámaras RTSP",
    "11. Selector de modo de detección",
    "12. Estadísticas del sistema",
    "13. Exportar incidentes a Excel",
    "14. Preguntas frecuentes y solución de problemas",
]
for item in indice:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h('1. Introducción y descripción del sistema')
p('CCTV AI PRO es un sistema de monitoreo inteligente desarrollado para la Unidad de '
  'Seguridad Ciudadana y el Serenazgo Municipal de la Municipalidad Distrital de Caracoto. '
  'Permite la vigilancia en tiempo real de múltiples cámaras RTSP con detección automática '
  'de tres tipos de incidentes críticos:')
for item in ['🔥 Incendios: detección por modelo YOLO entrenado + análisis de color HSV.',
             '🚗 Choques de vehículos: seguimiento de trayectorias y análisis de velocidad.',
             '🚨 Robos menores: detección de personas, objetos de valor y actividad sospechosa.']:
    doc.add_paragraph(item, style='List Bullet')
p('Cuando se detecta un incidente, el sistema lo registra en una base de datos local y '
  'puede enviar alertas automáticas por Telegram y/o Gmail a los operadores configurados.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. REQUISITOS
# ═══════════════════════════════════════════════════════════════════════════════
h('2. Requisitos mínimos del sistema')
t2 = doc.add_table(rows=7, cols=2); t2.style = 'Table Grid'
specs = [('Componente','Especificación'),
         ('Sistema operativo','Windows 10/11 (64-bit)'),
         ('CPU','Intel Core i5 o superior (8 hilos recomendados)'),
         ('GPU','NVIDIA con soporte CUDA (mínimo 2 GB VRAM)'),
         ('RAM','8 GB mínimo, 16 GB recomendado'),
         ('Almacenamiento','5 GB libres (modelos + base de datos)'),
         ('Red','LAN local con acceso a cámaras IP/RTSP')]
for i,(k,v) in enumerate(specs):
    t2.rows[i].cells[0].text = k; t2.rows[i].cells[1].text = v
    if i == 0:
        for j in range(2):
            t2.rows[i].cells[j].paragraphs[0].runs[0].bold = True
            t2.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd(t2.rows[i].cells[j], '2E74B5')
    elif i % 2 == 0:
        shd(t2.rows[i].cells[0], 'F0F5FB'); shd(t2.rows[i].cells[1], 'F0F5FB')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. INICIO
# ═══════════════════════════════════════════════════════════════════════════════
h('3. Inicio de la aplicación')
p('Para iniciar CCTV AI PRO sigue estos pasos:')
step_table([
    (1, 'Abre una ventana de PowerShell o Símbolo del sistema en la carpeta del proyecto.'),
    (2, 'Activa el entorno virtual ejecutando:  .\\practicas\\Scripts\\Activate.ps1'),
    (3, 'Ejecuta la aplicación con el comando:  python main.py'),
    (4, 'Espera a que se carguen los modelos de IA (puede tardar 15-30 segundos la primera vez).'),
    (5, 'La ventana principal se abrirá mostrando la vista de cámaras.'),
])
note('Si ves el mensaje "[BOOT] Optimizaciones de CPU aplicadas", la aplicación inició correctamente.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. INTERFAZ PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════
h('4. Interfaz principal — Vista de Cámaras')
p('La interfaz principal tiene tres zonas:')
for item in [
    'Barra superior (Top Bar): contiene el logo, el estado del sistema, el selector de modo y el botón de gestión de cámaras.',
    'Menú lateral izquierdo (Sidebar): permite navegar entre las diferentes vistas del sistema.',
    'Área de contenido central: cambia según la vista seleccionada en el menú lateral.',
]:
    doc.add_paragraph(item, style='List Bullet')
p('')
p('Vista de Cámaras — Elementos visibles:', bold=True)
for item in [
    '🟢 SYSTEM ONLINE: indica que el sistema está funcionando correctamente.',
    'Grilla de cámaras: muestra hasta 6 cámaras en formato 3×2. Las cámaras inactivas muestran "SIN VIDEO".',
    'Estadísticas del Sistema: panel inferior con frames procesados, modo activo, FPS y estado de cámaras.',
    'Selector de Modo: cambia entre Detección General, Accidentes, Personas o Robos.',
]:
    doc.add_paragraph(item, style='List Bullet')
note('La ventana es redimensionable. Puedes arrastrar los bordes para ajustar el tamaño según tu monitor.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MENÚ LATERAL
# ═══════════════════════════════════════════════════════════════════════════════
h('5. Menú lateral — Navegación')
p('El menú lateral izquierdo contiene 6 secciones. Haz clic en cada ítem para navegar:')
t5 = doc.add_table(rows=7, cols=2); t5.style = 'Table Grid'
menu_items = [
    ('Ítem del menú', 'Descripción'),
    ('Dashboard', 'Vista principal con grilla de cámaras y estadísticas del sistema.'),
    ('Cámaras', 'Igual que Dashboard. Vista de cámaras en tiempo real.'),
    ('Incidentes', 'Tabla con todos los incidentes detectados (incendios, choques, robos). Permite exportar a Excel.'),
    ('Destinatarios', 'Gestión de personas que recibirán las alertas automáticas.'),
    ('Notificaciones', 'Activar/desactivar canales de alerta: Telegram y Gmail.'),
    ('Ajustes', 'Controlar qué módulos están activos por cámara y ajustar el FPS objetivo.'),
]
for i,(k,v) in enumerate(menu_items):
    t5.rows[i].cells[0].text = k; t5.rows[i].cells[1].text = v
    if i == 0:
        for j in range(2):
            t5.rows[i].cells[j].paragraphs[0].runs[0].bold = True
            t5.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd(t5.rows[i].cells[j], '2E74B5')
    elif i % 2 == 0:
        shd(t5.rows[i].cells[0], 'F0F5FB'); shd(t5.rows[i].cells[1], 'F0F5FB')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. INCIDENTES
# ═══════════════════════════════════════════════════════════════════════════════
h('6. Vista de Incidentes')
p('La vista de Incidentes muestra un historial de los últimos 100 eventos detectados por el sistema.')
p('Columnas de la tabla:', bold=True)
for item in ['Hora: fecha y hora en que se detectó el incidente.',
             'Cámara: nombre de la cámara donde ocurrió.',
             'Tipo: fire (incendio), accident (choque) o theft (robo).',
             'Severidad: ALTO, MEDIO o BAJO según la confianza del detector.',
             'Confianza: valor numérico entre 0.0 y 1.0 del modelo de IA.']:
    doc.add_paragraph(item, style='List Bullet')
p('Cómo exportar incidentes a Excel:', bold=True)
step_table([
    (1, 'Ve a Incidentes desde el menú lateral.'),
    (2, 'Haz clic en el botón azul "Exportar a Excel" en la parte superior.'),
    (3, 'Se abrirá un diálogo para elegir la ubicación y nombre del archivo.'),
    (4, 'Haz clic en Guardar. El archivo .xlsx se creará con todos los incidentes.'),
])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. DESTINATARIOS
# ═══════════════════════════════════════════════════════════════════════════════
h('7. Vista de Destinatarios')
p('Los destinatarios son las personas que recibirán alertas automáticas cuando se detecte un incendio.')
p('Agregar un destinatario:', bold=True)
step_table([
    (1, 'Ve a Destinatarios desde el menú lateral.'),
    (2, 'Ingresa el Nombre del destinatario en el primer campo.'),
    (3, 'Ingresa su correo electrónico (opcional).'),
    (4, 'Ingresa su número de teléfono en formato internacional, Ej: +51912345678 (opcional).'),
    (5, 'Haz clic en "Agregar destinatario". El registro aparecerá en la tabla.'),
])
p('Eliminar un destinatario:', bold=True)
step_table([
    (1, 'Haz clic en la fila del destinatario que deseas eliminar para seleccionarla.'),
    (2, 'Haz clic en el botón rojo "Eliminar seleccionado".'),
    (3, 'El destinatario se eliminará permanentemente de la base de datos.'),
])
note('Los destinatarios se guardan en la base de datos local incidents.db. No se perderán al reiniciar.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. NOTIFICACIONES
# ═══════════════════════════════════════════════════════════════════════════════
h('8. Vista de Notificaciones — Activar Telegram y Gmail')
p('Desde esta vista puedes configurar qué canales de comunicación usará el sistema para '
  'enviar alertas automáticas cuando se detecte un incidente.')

h('8.1. Configurar Telegram', level=2, color=AZUL2)
step_table([
    (1, 'Ve a Notificaciones desde el menú lateral.'),
    (2, 'En la sección "Telegram", marca la casilla "Habilitado" para activar las alertas por Telegram.'),
    (3, 'Ingresa el Token del bot en el campo correspondiente (obtenido desde @BotFather en Telegram).'),
    (4, 'Ingresa los Chat ID(s) de los chats o grupos que recibirán las alertas (separados por coma).'),
    (5, 'Haz clic en "💾 Guardar configuración de notificaciones".'),
    (6, 'Para desactivar Telegram: desmarca la casilla "Habilitado" y guarda.'),
])
note('Para obtener tu Chat ID: escribe /start a tu bot y visita: https://api.telegram.org/botTU_TOKEN/getUpdates')

h('8.2. Configurar Gmail', level=2, color=AZUL2)
step_table([
    (1, 'En la sección "Gmail", marca la casilla "Habilitado".'),
    (2, 'Ingresa el correo remitente (tu cuenta de Gmail).'),
    (3, 'Ingresa la Contraseña de aplicación (generada en myaccount.google.com > Seguridad > Contraseñas de app).'),
    (4, 'Ingresa los correos destinatarios separados por coma.'),
    (5, 'Haz clic en "💾 Guardar configuración de notificaciones".'),
    (6, 'Para desactivar Gmail: desmarca la casilla "Habilitado" y guarda.'),
])
note('IMPORTANTE: Usa una Contraseña de aplicación, no tu contraseña normal de Gmail. Las contraseñas de aplicación tienen 16 caracteres sin espacios.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. AJUSTES
# ═══════════════════════════════════════════════════════════════════════════════
h('9. Vista de Ajustes — Módulos por cámara y FPS')
p('La vista de Ajustes te permite controlar qué módulos de IA están activos en cada cámara '
  'y ajustar el FPS objetivo para gestionar el consumo de recursos del sistema.')

h('9.1. Habilitar/Deshabilitar módulos por cámara', level=2, color=AZUL2)
p('Cada cámara configurada aparece como un panel con tres casillas de verificación:')
for item in [
    '🔥 Incendios: activa/desactiva la detección de fuego y humo en esa cámara.',
    '🚨 Robos: activa/desactiva el análisis de personas y objetos sospechosos.',
    '🚗 Choques: activa/desactiva el seguimiento de vehículos y detección de colisiones.',
]:
    doc.add_paragraph(item, style='List Bullet')
step_table([
    (1, 'Ve a Ajustes desde el menú lateral.'),
    (2, 'Localiza el panel de la cámara que deseas configurar.'),
    (3, 'Marca o desmarca las casillas de los módulos según lo que necesites.'),
    (4, 'Haz clic en "💾 Guardar ajustes". Los cambios se aplican inmediatamente sin reiniciar.'),
])
note('Recomendación: si el sistema va lento, desactiva el módulo de Robos en las cámaras donde no sea necesario, ya que es el que más recursos consume.')

h('9.2. Ajustar el FPS objetivo', level=2, color=AZUL2)
p('El campo "FPS objetivo" controla cuántos frames por segundo procesa el sistema. '
  'Valores más bajos reducen el consumo de CPU/GPU.')
t9 = doc.add_table(rows=4, cols=3); t9.style = 'Table Grid'
fps_data = [('FPS','Uso de recursos','Caso de uso recomendado'),
            ('5-10','Bajo (CPU < 30%)','Edificios, estacionamientos (movimiento lento)'),
            ('15-20','Medio (CPU 40-60%)','Calles, pasillos (uso normal)'),
            ('25-30','Alto (GPU recomendada)','Intersecciones viales, zonas de alta actividad')]
for i,(a,b,c_) in enumerate(fps_data):
    t9.rows[i].cells[0].text = a; t9.rows[i].cells[1].text = b; t9.rows[i].cells[2].text = c_
    if i == 0:
        for j in range(3):
            t9.rows[i].cells[j].paragraphs[0].runs[0].bold = True
            t9.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd(t9.rows[i].cells[j], '2E74B5')
    elif i % 2 == 0:
        for j in range(3): shd(t9.rows[i].cells[j], 'F0F5FB')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. GESTIÓN DE CÁMARAS
# ═══════════════════════════════════════════════════════════════════════════════
h('10. Gestión de cámaras RTSP')
p('Para agregar, editar o eliminar cámaras haz clic en el botón "📹 Gestionar Cámaras" de la barra superior.')
step_table([
    (1, 'Haz clic en "📹 Gestionar Cámaras" en la barra superior.'),
    (2, 'Se abrirá el diálogo de gestión de cámaras.'),
    (3, 'Para agregar una cámara: ingresa el nombre y la URL RTSP y haz clic en "Agregar".'),
    (4, 'Formato URL RTSP típico:  rtsp://usuario:contraseña@192.168.1.100:554/stream'),
    (5, 'Para eliminar una cámara: selecciónala en la lista y haz clic en "Eliminar".'),
    (6, 'Cierra el diálogo. Los cambios se guardan automáticamente.'),
])
note('Las cámaras configuradas se guardan en cameras_config.json. Al reiniciar la aplicación se cargarán automáticamente.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. SELECTOR DE MODO
# ═══════════════════════════════════════════════════════════════════════════════
h('11. Selector de modo de detección')
p('El selector de modo en la barra superior define qué tipo de análisis se aplica sobre los frames de video.')
t11 = doc.add_table(rows=5, cols=2); t11.style = 'Table Grid'
modos = [('Modo','Descripción'),
         ('🎯 Detección General','Usa el modelo YOLOv8 base (yolov8n.pt) para detectar cualquier objeto. Sin análisis especializado.'),
         ('🚗 Detección de Accidentes','Activa el seguimiento de vehículos. Detecta colisiones por proximidad y desaceleración brusca.'),
         ('👤 Rastreo de Personas','Identifica y sigue personas usando histograma de color. ID estable entre frames.'),
         ('🚨 Detección de Robos','Detecta personas con objetos de valor y analiza transferencias sospechosas. El más intensivo.')]
for i,(k,v) in enumerate(modos):
    t11.rows[i].cells[0].text = k; t11.rows[i].cells[1].text = v
    if i == 0:
        for j in range(2):
            t11.rows[i].cells[j].paragraphs[0].runs[0].bold = True
            t11.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd(t11.rows[i].cells[j], '2E74B5')
    elif i % 2 == 0:
        shd(t11.rows[i].cells[0], 'F0F5FB'); shd(t11.rows[i].cells[1], 'F0F5FB')
note('La detección de Incendios se ejecuta siempre en todas las cámaras, independientemente del modo seleccionado.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. ESTADÍSTICAS
# ═══════════════════════════════════════════════════════════════════════════════
h('12. Estadísticas del sistema')
p('El panel de Estadísticas aparece únicamente en la vista de Cámaras (Dashboard), en la parte inferior. '
  'Se actualiza cada 60 frames (aproximadamente cada 2.4 segundos a 25 FPS). Muestra:')
for item in [
    'Frames procesados: total de frames analizados desde que inició la sesión.',
    'Modo: modo de detección activo actualmente.',
    'GPU: hardware de procesamiento detectado.',
    'FPS objetivo: velocidad de procesamiento configurada.',
    'Cámaras configuradas / activas: número de cámaras con señal de video.',
    'Resolución: tamaño de frame que se envía al modelo de IA.',
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. EXPORTAR A EXCEL
# ═══════════════════════════════════════════════════════════════════════════════
h('13. Exportar incidentes a Excel')
step_table([
    (1, 'Haz clic en "Incidentes" en el menú lateral.'),
    (2, 'Haz clic en el botón azul "Exportar a Excel".'),
    (3, 'Elige la carpeta de destino y el nombre del archivo (por defecto: incidents.xlsx).'),
    (4, 'Haz clic en Guardar. Se generará un archivo Excel con todas las columnas.'),
    (5, 'Abre el archivo con Microsoft Excel o LibreOffice Calc.'),
])
note('El archivo Excel incluye: ID, Hora, Cámara, Tipo de incidente, Severidad, Confianza y Detalles.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════════════════════════════
h('14. Preguntas frecuentes y solución de problemas')

problemas = [
    ('¿Por qué la cámara muestra "SIN VIDEO"?',
     'Verifica que: (1) la URL RTSP sea correcta, (2) la cámara esté encendida y en la misma red, '
     '(3) el usuario y contraseña en la URL sean correctos. Usa "📹 Gestionar Cámaras" para editar la URL.'),
    ('La aplicación inicia lento (20-30 segundos)',
     'Es normal. La primera vez debe cargar 3 modelos YOLOv8 a la GPU. Las siguientes veces '
     'puede ser más rápido si los modelos están en caché de CUDA.'),
    ('No llegan alertas de Telegram',
     'Verifica en Notificaciones que: (1) "Habilitado" esté marcado, (2) el Token del bot sea correcto, '
     '(3) hayas enviado /start al bot desde tu cuenta de Telegram antes de recibir mensajes.'),
    ('No llegan alertas de Gmail',
     'Verifica: (1) que uses una Contraseña de aplicación (no tu contraseña normal), '
     '(2) que la cuenta Gmail tenga activada la verificación en 2 pasos, '
     '(3) que el correo remitente coincida con la cuenta configurada.'),
    ('El sistema va lento o consume mucho CPU',
     'Ve a Ajustes y: (1) reduce el FPS objetivo a 10-15, (2) desactiva el módulo "Robos" en cámaras '
     'donde no sea necesario (es el más intensivo), (3) desactiva el módulo de "Choques" en cámaras interiores.'),
    ('¿Cómo sé si un incidente fue detectado correctamente?',
     'Ve a "Incidentes" en el menú lateral. Si el incidente aparece en la tabla con una confianza > 0.5, '
     'fue una detección confiable. Confianza < 0.3 puede ser un falso positivo.'),
    ('¿Dónde se guardan los datos del sistema?',
     'incidents.db (base de datos SQLite), cameras_config.json (cámaras), '
     'notifications_config.json (notificaciones), fire_config.json (umbrales de incendio).'),
]

for pregunta, respuesta in problemas:
    p(pregunta, bold=True)
    p(respuesta)
    doc.add_paragraph()

# Firma
doc.add_page_break()
c('Juliaca, julio del 2026')
doc.add_paragraph()
c('_________________________________')
c('Yords Williams Ccalla Mamani', bold=True)
c('Practicante — Municipalidad Distrital de Caracoto')
c('Escuela Profesional de Ingeniería de Software y Sistemas — UNAJ')

out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'MANUAL_USUARIO.docx')
doc.save(out)
print(f'Guardado: {out}')
