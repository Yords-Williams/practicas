"""
Genera MANUAL_USUARIO_v2.docx — versión mejorada con diagramas, mockups y mejor diseño.
"""
import os, io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── Paleta de colores ─────────────────────────────────────────────────────────
AZUL    = RGBColor(0x1F, 0x2A, 0x5E)
AZUL2   = RGBColor(0x2E, 0x74, 0xB5)
VERDE   = RGBColor(0x1E, 0x8B, 0x4C)
ROJO    = RGBColor(0xC0, 0x39, 0x2B)
NARANJA = RGBColor(0xD4, 0x6B, 0x08)
GRIS    = RGBColor(0x55, 0x55, 0x55)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3);  s.right_margin=Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shd(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    e=OxmlElement('w:shd')
    e.set(qn('w:val'),'clear'); e.set(qn('w:color'),'auto')
    e.set(qn('w:fill'),hex_color); tcPr.append(e)

def h(text, level=1, color=AZUL):
    ph=doc.add_heading(text, level=level)
    ph.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in ph.runs: r.font.color.rgb=color; r.bold=True
    return ph

def p(text, size=11, bold=False, italic=False, color=None, justify=True):
    para=doc.add_paragraph()
    if justify: para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=para.add_run(text); r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    return para

def c(text, bold=False, size=12, color=None):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(text); r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return para

def img_from_fig(fig, width_cm=14, caption=None):
    buf=io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight', dpi=130)
    buf.seek(0); plt.close(fig)
    doc.add_picture(buf, width=Cm(width_cm))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size=Pt(9); cp.runs[0].italic=True
        cp.runs[0].font.color.rgb=GRIS

def img_file(path, width_cm=13, caption=None):
    if path and os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size=Pt(9); cp.runs[0].italic=True
        cp.runs[0].font.color.rgb=GRIS

def callout(text, tipo='nota'):
    """Caja de aviso: nota (azul), advertencia (naranja), peligro (rojo)."""
    colores={'nota':('2E74B5','DCE6F1','NOTA'),'advertencia':('D4340A','FDE9E0','ADVERTENCIA'),
             'consejo':('1E8B4C','E8F5E9','CONSEJO')}
    hex_border, hex_bg, label = colores.get(tipo, colores['nota'])
    t=doc.add_table(rows=1, cols=1); t.style='Table Grid'
    cell=t.rows[0].cells[0]
    shd(cell, hex_bg)
    para=cell.paragraphs[0]; para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r1=para.add_run(f'{label}: ')
    r1.bold=True; r1.font.size=Pt(10)
    r1.font.color.rgb=RGBColor.from_string(hex_border)
    r2=para.add_run(text); r2.font.size=Pt(10)
    doc.add_paragraph()

def steps_table(steps):
    t=doc.add_table(rows=len(steps), cols=2); t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(num,text) in enumerate(steps):
        t.rows[i].cells[0].text=str(num)
        t.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        t.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=BLANCO
        shd(t.rows[i].cells[0], '2E74B5')
        t.rows[i].cells[0].width=Cm(1.2)
        t.rows[i].cells[1].text=text
        t.rows[i].cells[1].paragraphs[0].runs[0].font.size=Pt(10.5)
        if i%2==0: shd(t.rows[i].cells[1],'F0F5FB')
    doc.add_paragraph()

def feature_table(headers, rows_data):
    t=doc.add_table(rows=1+len(rows_data), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hdr in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=hdr
        c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.color.rgb=BLANCO
        shd(c,'1F2A5E')
    for i,row in enumerate(rows_data):
        for j,val in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=val
            c.paragraphs[0].runs[0].font.size=Pt(10)
            if i%2==0: shd(c,'F0F5FB')
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Vista principal de la aplicación
# ══════════════════════════════════════════════════════════════════════════════
def make_ui_mockup_cameras():
    fig,ax=plt.subplots(figsize=(13,7.5))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7.5); ax.axis('off')

    def rect(x,y,w,h,fc,ec='none',alpha=1.0,radius=0.15):
        r=FancyBboxPatch((x,y),w,h,boxstyle=f"round,pad={radius}",
                          facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha)
        ax.add_patch(r)
    def txt(x,y,t,size=8,color='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=size,color=color,
                ha=ha,va=va,fontweight='bold' if bold else 'normal')

    # Sidebar
    rect(0,0,2,7.5,'#0b0d0f')
    txt(1,7.0,'CCTV AI PRO',9,'#e6eef6',bold=True)
    menu_items=[('Dashboard','#ffb3b3',6.3),('Camaras','#e6eef6',5.7),
                ('Incidentes','#e6eef6',5.1),('Destinatarios','#e6eef6',4.5),
                ('Notificaciones','#e6eef6',3.9),('Ajustes','#e6eef6',3.3)]
    for label,color,ypos in menu_items:
        if color=='#ffb3b3':
            rect(0.15,ypos-0.2,1.7,0.4,'#2b2b2b',radius=0.05)
        txt(1,ypos,label,8.5,color)

    # Topbar
    rect(2,6.9,11,0.6,'#1f2a36')
    rect(2.2,6.95,1.2,0.45,'#27ae60',radius=0.08)
    txt(2.8,7.18,'SYSTEM ONLINE',7,'white',bold=True)
    rect(4.2,6.97,2.5,0.4,'#22313f',ec='#3b3b3b',radius=0.06)
    txt(5.45,7.17,'Deteccion General',7.5,'white')
    rect(7.2,6.97,1.8,0.4,'#e74c3c',radius=0.08)
    txt(8.1,7.17,'Gestionar Camaras',7,'white',bold=True)

    # Grid de cámaras (3×2)
    cam_pos=[(2.1,3.9),(5.55,3.9),(9.0,3.9),(2.1,0.3),(5.55,0.3),(9.0,0.3)]
    cam_labels=['Camara 1\nEntrada','Camara 2\nParqueo','SIN VIDEO',
                'SIN VIDEO','SIN VIDEO','SIN VIDEO']
    cam_colors=['#0f1113','#0f1113','#1a1a1a','#1a1a1a','#1a1a1a','#1a1a1a']
    for (x,y),label,fc in zip(cam_pos,cam_labels,cam_colors):
        rect(x,y,3.1,3.4,fc,'#2b2b2b',radius=0.12)
        if 'SIN VIDEO' in label:
            txt(x+1.55,y+1.7,label,9,'#7f8c8d')
        else:
            # Simular imagen de cámara
            rect(x+0.08,y+0.08,2.95,3.25,'#1a2535',radius=0.08)
            # Grid lines para simular imagen
            for gx in np.linspace(x+0.3,x+2.8,4):
                ax.plot([gx,gx],[y+0.15,y+3.25],'--',color='#2a3a4a',lw=0.5,alpha=0.4)
            for gy in np.linspace(y+0.3,y+3.1,4):
                ax.plot([x+0.15,x+2.95],[gy,gy],'--',color='#2a3a4a',lw=0.5,alpha=0.4)
            # Bounding box simulado (detección)
            rect(x+0.6,y+0.8,0.8,1.2,'none','#00ff88',alpha=0.9,radius=0.04)
            txt(x+1.0,y+2.1,'Persona 87%',5.5,'#00ff88')
            txt(x+1.55,y+3.0,label,7,'#bfc9d4')

    # Panel de estadísticas (abajo)
    rect(2.1,0.02,10.7,0.26,'#111214','#222',radius=0.04)
    stats='Frames: 1247  |  Modo: Deteccion General  |  GPU: Quadro P1000  |  FPS: 25  |  Camaras: 2/6'
    txt(7.45,0.15,stats,6.5,'#dfe7ec')

    ax.set_title('Vista de Camaras — CCTV AI PRO', fontsize=11,
                 fontweight='bold', color='#e6eef6', pad=6)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Vista de Notificaciones
# ══════════════════════════════════════════════════════════════════════════════
def make_ui_mockup_notif():
    fig,ax=plt.subplots(figsize=(13,7))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis('off')

    def rect(x,y,w,h,fc,ec='none',alpha=1.0,radius=0.1):
        r=FancyBboxPatch((x,y),w,h,boxstyle=f"round,pad={radius}",
                          facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha)
        ax.add_patch(r)
    def txt(x,y,t,size=9,color='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=size,color=color,ha=ha,va=va,
                fontweight='bold' if bold else 'normal')
    def field(x,y,w,label,value='',enabled=None):
        txt(x,y+0.25,label,8,'#bfc9d4',ha='left')
        rect(x,y-0.05,w,0.35,'#1a1a2e','#3b3b3b',radius=0.06)
        if value: txt(x+0.1,y+0.13,value,8,'#e6eef6',ha='left')
    def checkbox(x,y,checked=True,label=''):
        fc='#2e74b5' if checked else '#1a1a2e'
        rect(x,y,0.3,0.28,fc,'#3b3b3b',radius=0.04)
        if checked: txt(x+0.15,y+0.14,'✓',9,'white',bold=True)
        txt(x+0.4,y+0.14,label,9,'#e6eef6',ha='left')

    # Sidebar
    rect(0,0,2,7,'#0b0d0f')
    txt(1,6.6,'CCTV AI PRO',9,'#e6eef6',bold=True)
    for label,color,ypos in [('Notificaciones','#ffb3b3',4.2),('Ajustes','#e6eef6',3.6)]:
        if color=='#ffb3b3': rect(0.1,ypos-0.18,1.8,0.38,'#2b2b2b',radius=0.05)
        txt(1,ypos,label,8.5,color)

    # Topbar
    rect(2,6.4,11,0.55,'#1f2a36')
    txt(7,6.67,'Configuracion de Notificaciones',10,'#e6eef6',bold=True)

    # Panel principal
    txt(2.3,6.05,'Configura los canales por los que se envian alertas de incidentes.',
        8.5,'#99aab5',ha='left')

    # Sección Telegram
    rect(2.1,3.4,10.5,2.4,'#111214','#2e74b5',radius=0.12)
    txt(3.0,5.6,'TELEGRAM',10,'#2e74b5',bold=True,ha='left')
    checkbox(2.3,5.05,checked=True,label='Habilitado')
    field(2.3,4.4,4.5,'Token del bot:','••••••••••••••••••••••')
    field(7.2,4.4,5.0,'Chat ID(s):','7973977029')
    field(2.3,3.6,4.5,'Estado:','Conectado - 1 chat')

    # Sección Gmail
    rect(2.1,0.5,10.5,2.6,'#111214','#27ae60',radius=0.12)
    txt(3.0,2.85,'GMAIL',10,'#27ae60',bold=True,ha='left')
    checkbox(2.3,2.3,checked=True,label='Habilitado')
    field(2.3,1.65,4.5,'Remitente:','yordswcm@gmail.com')
    field(7.2,1.65,5.0,'Contrasena de app:','•••• •••• •••• ••••')
    field(2.3,0.85,9.9,'Destinatarios:','closbtep@gmail.com')

    # Botón guardar
    rect(8.5,0.1,3.8,0.35,'#27ae60',radius=0.08)
    txt(10.4,0.28,'Guardar configuracion',9,'white',bold=True)

    ax.set_title('Vista de Notificaciones — Activar/Desactivar canales', fontsize=11,
                 fontweight='bold', color='#e6eef6', pad=6)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Vista de Ajustes
# ══════════════════════════════════════════════════════════════════════════════
def make_ui_mockup_ajustes():
    fig,ax=plt.subplots(figsize=(13,7))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,7); ax.axis('off')

    def rect(x,y,w,h,fc,ec='none',alpha=1.0,radius=0.1):
        r=FancyBboxPatch((x,y),w,h,boxstyle=f"round,pad={radius}",
                         facecolor=fc,edgecolor=ec,linewidth=1.5,alpha=alpha)
        ax.add_patch(r)
    def txt(x,y,t,size=9,color='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=size,color=color,ha=ha,va=va,
                fontweight='bold' if bold else 'normal')
    def checkbox(x,y,checked=True,label='',lcolor='#e6eef6'):
        fc='#2e74b5' if checked else '#1a1a2e'
        rect(x,y,0.28,0.26,fc,'#3b3b3b',radius=0.04)
        if checked: txt(x+0.14,y+0.13,'✓',8,'white',bold=True)
        txt(x+0.38,y+0.13,label,9,lcolor,ha='left')

    # Sidebar
    rect(0,0,2,7,'#0b0d0f')
    txt(1,6.6,'CCTV AI PRO',9,'#e6eef6',bold=True)
    for label,color,ypos in [('Ajustes','#ffb3b3',3.9),('Notificaciones','#e6eef6',4.5)]:
        if color=='#ffb3b3': rect(0.1,ypos-0.18,1.8,0.38,'#2b2b2b',radius=0.05)
        txt(1,ypos,label,8.5,color)

    # Topbar
    rect(2,6.4,11,0.55,'#1f2a36')
    txt(7,6.67,'Ajustes del Sistema',10,'#e6eef6',bold=True)

    txt(2.3,6.05,'Controla los modulos activos por camara y el FPS objetivo.',
        8.5,'#99aab5',ha='left')

    # Cámara 1
    rect(2.1,4.1,10.5,1.65,'#111214','#333',radius=0.12)
    txt(3.0,5.55,'Camara 1 — Entrada Principal',9.5,'#e6eef6',bold=True,ha='left')
    checkbox(2.4,4.85,True,'Incendios','#e67e22')
    checkbox(5.3,4.85,True,'Robos','#8e44ad')
    checkbox(8.2,4.85,True,'Choques','#c0392b')
    txt(2.4,4.4,'Todos los modulos activos (mayor precision)',8,'#27ae60',ha='left')

    # Cámara 2
    rect(2.1,2.2,10.5,1.65,'#111214','#333',radius=0.12)
    txt(3.0,3.65,'Camara 2 — Parqueo Exterior',9.5,'#e6eef6',bold=True,ha='left')
    checkbox(2.4,2.95,True,'Incendios','#e67e22')
    checkbox(5.3,2.95,False,'Robos','#666')
    checkbox(8.2,2.95,True,'Choques','#c0392b')
    txt(2.4,2.5,'Robos desactivado (menor consumo de GPU)',8,'#e67e22',ha='left')

    # FPS
    rect(2.1,0.5,10.5,1.45,'#111214','#333',radius=0.12)
    txt(3.0,1.72,'Rendimiento',9.5,'#e6eef6',bold=True,ha='left')
    txt(2.4,1.25,'FPS objetivo (1-30):',8.5,'#bfc9d4',ha='left')
    rect(6.5,1.0,1.2,0.4,'#1a1a2e','#3b3b3b',radius=0.06)
    txt(7.1,1.2,'25',10,'#e6eef6',bold=True)
    for val,xpos in [(5,3.5),(10,5.0),(15,6.5),(20,8.0),(25,9.5),(30,11.0)]:
        ax.plot(xpos+2.1*0,1.35,'|',color='#555',ms=6)
        txt(xpos*0.35+3.8,0.72,str(val),7,'#99aab5')
    rect(2.4,0.6,9.9*0.85,0.22,'#333',radius=0.04)
    rect(2.4,0.6,9.9*0.85*0.83,0.22,'#2e74b5',radius=0.04)

    # Botón guardar
    rect(8.5,0.12,3.8,0.35,'#2980b9',radius=0.08)
    txt(10.4,0.3,'Guardar ajustes',9,'white',bold=True)

    ax.set_title('Vista de Ajustes — Modulos por camara y FPS', fontsize=11,
                 fontweight='bold', color='#e6eef6', pad=6)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Vista de Incidentes
# ══════════════════════════════════════════════════════════════════════════════
def make_ui_mockup_incidents():
    fig,ax=plt.subplots(figsize=(13,6))
    fig.patch.set_facecolor('#0b1220'); ax.set_facecolor('#0b1220')
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis('off')

    def rect(x,y,w,h,fc,ec='none',alpha=1.0,radius=0.08):
        r=FancyBboxPatch((x,y),w,h,boxstyle=f"round,pad={radius}",
                         facecolor=fc,edgecolor=ec,linewidth=1,alpha=alpha)
        ax.add_patch(r)
    def txt(x,y,t,size=8.5,color='white',bold=False,ha='center',va='center'):
        ax.text(x,y,t,fontsize=size,color=color,ha=ha,va=va,
                fontweight='bold' if bold else 'normal')

    # Sidebar
    rect(0,0,2,6,'#0b0d0f')
    txt(1,5.6,'CCTV AI PRO',9,'#e6eef6',bold=True)
    for label,color,ypos in [('Incidentes','#ffb3b3',3.3)]:
        rect(0.1,ypos-0.18,1.8,0.38,'#2b2b2b',radius=0.05)
        txt(1,ypos,label,9,color)

    # Topbar
    rect(2,5.4,11,0.55,'#1f2a36')
    txt(7,5.67,'Incidentes Detectados',10,'#e6eef6',bold=True)

    # Barra de herramientas
    rect(2.1,4.9,2.0,0.4,'#2980b9',radius=0.07)
    txt(3.1,5.1,'Exportar a Excel',8,'white',bold=True)

    # Encabezado de tabla
    headers=['Hora','Camara','Tipo','Severidad','Confianza']
    widths=[2.0,2.5,2.5,2.0,1.5]
    x_start=2.1
    for hdr,w in zip(headers,widths):
        rect(x_start,4.35,w-0.05,0.45,'#1f2a5e',radius=0.0)
        txt(x_start+w/2-0.025,4.58,hdr,8.5,'#e6eef6',bold=True)
        x_start+=w

    # Filas de la tabla
    incidents=[
        ('2026-07-15 14:23:11','Camara 1','fire','ALTO','0.91'),
        ('2026-07-15 13:55:04','Camara 2','accident','ALTO','0.78'),
        ('2026-07-15 13:22:47','Camara 1','theft','MEDIO','0.72'),
        ('2026-07-15 12:14:33','Camara 1','fire','MEDIO','0.58'),
        ('2026-07-15 11:09:18','Camara 2','accident','BAJO','0.43'),
    ]
    sev_colors={'ALTO':'#e74c3c','MEDIO':'#e67e22','BAJO':'#27ae60'}
    tipo_icons={'fire':'Incendio','accident':'Choque','theft':'Robo'}
    for ri,(hora,cam,tipo,sev,conf) in enumerate(incidents):
        y=4.35-(ri+1)*0.5
        bg='#111214' if ri%2==0 else '#0f1113'
        x_start=2.1
        for val,w in zip([hora,cam,tipo_icons[tipo],sev,conf],widths):
            rect(x_start,y,w-0.05,0.44,bg,radius=0.0)
            color_v='#e6eef6'
            if val in sev_colors: color_v=sev_colors[val]
            elif val in ('Incendio','Choque','Robo'):
                icon_c={'Incendio':'#e67e22','Choque':'#c0392b','Robo':'#8e44ad'}
                color_v=icon_c[val]
            txt(x_start+w/2-0.025,y+0.22,val,8,color_v,bold=(val in sev_colors))
            x_start+=w

    ax.set_title('Vista de Incidentes — Historial con exportacion a Excel', fontsize=11,
                 fontweight='bold', color='#e6eef6', pad=6)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Modos de detección
# ══════════════════════════════════════════════════════════════════════════════
def make_detection_modes():
    fig,axes=plt.subplots(1,4,figsize=(14,3.5))
    fig.patch.set_facecolor('#111')
    modes=[
        ('General','#2E74B5','yolov8n.pt\nDetecta cualquier\nobjeto del dataset\nCOCO (80 clases)'),
        ('Accidentes','#C0392B','choques/best.pt\nTrackeo de vehiculos\nAnalisis de velocidad\ny proximidad'),
        ('Personas','#27AE60','person_identifier.py\nRe-identificacion\npor histograma HSV\nID estable'),
        ('Robos','#8E44AD','yolov8n.pt + SORT\nDeteccion personas\n+ objetos de valor\nScoring de robo'),
    ]
    icons=['[General]','[Choques]','[Personas]','[Robos]']
    for ax,(mode,color,desc),icon in zip(axes,modes,icons):
        ax.set_facecolor('#1a1a2e')
        ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis('off')
        # Header
        rect_h=FancyBboxPatch((0,0.75),1,0.25,boxstyle="round,pad=0.02",
                               facecolor=color,edgecolor='none')
        ax.add_patch(rect_h)
        ax.text(0.5,0.875,mode,ha='center',va='center',fontsize=12,
                fontweight='bold',color='white')
        # Content
        ax.text(0.5,0.38,desc,ha='center',va='center',fontsize=8.5,
                color='#e6eef6',multialignment='center',linespacing=1.5)
        # Border
        rect_b=FancyBboxPatch((0.02,0.02),0.96,0.98,boxstyle="round,pad=0.02",
                               facecolor='none',edgecolor=color,linewidth=2)
        ax.add_patch(rect_b)
    fig.suptitle('Modos de Deteccion Disponibles', fontsize=12,
                 fontweight='bold', color='#e6eef6', y=1.02)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# MOCKUP: Gestión de cámaras
# ══════════════════════════════════════════════════════════════════════════════
def make_camera_dialog():
    fig,ax=plt.subplots(figsize=(10,5))
    fig.patch.set_facecolor('#1a1a2e'); ax.set_facecolor('#1a1a2e')
    ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')

    def rect(x,y,w,h,fc,ec='none',radius=0.1):
        r=FancyBboxPatch((x,y),w,h,boxstyle=f"round,pad={radius}",
                         facecolor=fc,edgecolor=ec,linewidth=1.5)
        ax.add_patch(r)
    def txt(x,y,t,size=9,color='white',bold=False,ha='center'):
        ax.text(x,y,t,fontsize=size,color=color,ha=ha,va='center',
                fontweight='bold' if bold else 'normal')

    # Fondo principal
    rect(0.2,0.2,9.6,4.6,'#111314','#2e74b5')
    txt(5,4.5,'Gestion de Camaras',11,'#e6eef6',bold=True)
    ax.plot([0.2,9.8],[4.2,4.2],'-',color='#2e74b5',lw=1.5)

    # Lista de cámaras
    txt(0.8,3.9,'Camaras configuradas:',9,'#bfc9d4',ha='left')
    for i,(name,url) in enumerate([
        ('Camara 1 - Entrada','rtsp://admin:pass@192.168.1.100:554/stream'),
        ('Camara 2 - Parqueo','rtsp://admin:pass@192.168.1.101:554/stream'),
    ]):
        bg='#1f2a36' if i==0 else '#111314'
        rect(0.5,3.3-i*0.65,5.8,0.55,bg,'#333',radius=0.06)
        txt(0.7,3.58-i*0.65,name,8.5,'#e6eef6',ha='left')
        txt(0.7,3.42-i*0.65,url,7,'#7f8c8d',ha='left')

    # Formulario agregar
    txt(0.8,1.85,'Agregar nueva camara:',9,'#bfc9d4',ha='left')
    rect(0.5,1.3,2.5,0.45,'#1a1a2e','#3b3b3b',radius=0.06)
    txt(0.7,1.53,'Nombre de la camara',7.5,'#555',ha='left')
    rect(3.2,1.3,5.6,0.45,'#1a1a2e','#3b3b3b',radius=0.06)
    txt(3.4,1.53,'rtsp://usuario:contrasena@IP:554/stream',7.5,'#555',ha='left')

    # Botones
    rect(0.5,0.5,2,0.45,'#27ae60',radius=0.08)
    txt(1.5,0.73,'+ Agregar',9,'white',bold=True)
    rect(3.0,0.5,2.5,0.45,'#e74c3c',radius=0.08)
    txt(4.25,0.73,'Eliminar seleccionado',8,'white',bold=True)
    rect(6.2,0.5,3.3,0.45,'#2980b9',radius=0.08)
    txt(7.85,0.73,'Probar conexion',8.5,'white',bold=True)

    ax.set_title('Dialogo de Gestion de Camaras RTSP', fontsize=11,
                 fontweight='bold', color='#e6eef6', pad=8)
    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO WORD
# ══════════════════════════════════════════════════════════════════════════════

# ── PORTADA ───────────────────────────────────────────────────────────────────
logo_path=os.path.join(ROOT,'assets','app.png')
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Cm(3.5))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
c('UNIVERSIDAD NACIONAL DE JULIACA', bold=True, size=14, color=AZUL)
c('Facultad de Ciencias de Ingenierías', bold=True, size=11, color=AZUL)
c('E.P. Ingeniería de Software y Sistemas', size=10, color=AZUL)
doc.add_paragraph()
c('MANUAL DE USUARIO', bold=True, size=18, color=AZUL)
c('CCTV AI PRO', bold=True, size=14, color=AZUL2)
c('Sistema de Monitoreo Inteligente con Detección de Incidentes', size=11, color=GRIS)
doc.add_paragraph()
t=doc.add_table(rows=5,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([
    ('Versión:','2.0 — Julio 2026'),
    ('Sistema:','CCTV AI PRO'),
    ('Institución:','Municipalidad Distrital de Caracoto'),
    ('Elaborado por:','Yords Williams Ccalla Mamani'),
    ('Área de uso:','Seguridad Ciudadana — Serenazgo Municipal'),
]):
    t.rows[i].cells[0].text=k; t.rows[i].cells[1].text=v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold=True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size=Pt(10.5)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size=Pt(10.5)
    t.rows[i].cells[0].width=Cm(5); t.rows[i].cells[1].width=Cm(10)
doc.add_page_break()

# ── INTRODUCCIÓN ─────────────────────────────────────────────────────────────
h('Introducción')
p('CCTV AI PRO es un sistema de monitoreo inteligente que analiza automáticamente las '
  'imágenes de sus cámaras de seguridad para detectar incendios, robos menores y choques '
  'de vehículos en tiempo real. Cuando detecta un incidente, lo registra en la base de '
  'datos y puede enviarte una alerta por Telegram o Gmail de forma inmediata.')
p('Este manual te guía paso a paso en el uso de todas las funciones del sistema.')
doc.add_paragraph()
callout('Este manual está pensado para el personal operativo del Serenazgo Municipal. '
        'No se necesitan conocimientos técnicos para usar el sistema.', 'consejo')

# ── REQUISITOS ────────────────────────────────────────────────────────────────
h('1. Requisitos del Sistema')
feature_table(
    ['Componente','Mínimo requerido','Recomendado'],
    [('Windows','10 (64-bit)','11 (64-bit)'),
     ('GPU NVIDIA','1 GB VRAM (CUDA)','4+ GB VRAM'),
     ('RAM','8 GB','16 GB'),
     ('CPU','Intel Core i5 / 4 núcleos','Intel Core i7 / 8+ núcleos'),
     ('Red','LAN con acceso a cámaras IP','Gigabit Ethernet'),
     ('Almacenamiento','5 GB libres','10+ GB libres')]
)
callout('Si el equipo no tiene GPU NVIDIA, el sistema funcionará en modo CPU. '
        'Las detecciones serán más lentas pero igualmente funcionales.', 'advertencia')

# ── INICIO ────────────────────────────────────────────────────────────────────
h('2. Cómo Iniciar la Aplicación')
steps_table([
    (1, 'Abre una ventana de PowerShell en la carpeta del proyecto (clic derecho → "Abrir en Terminal").'),
    (2, 'Activa el entorno virtual ejecutando:   .\\practicas\\Scripts\\Activate.ps1'),
    (3, 'Si el sistema pregunta por permisos de ejecución, escribe:  Set-ExecutionPolicy RemoteSigned -Scope Process'),
    (4, 'Ejecuta la aplicación con:   python main.py'),
    (5, 'Espera mientras el sistema carga los modelos de IA. Verás mensajes como "[BOOT] Optimizaciones de CPU aplicadas".'),
    (6, 'La ventana principal se abrirá automáticamente. La carga tarda entre 15 y 30 segundos la primera vez.'),
])
callout('Si ves el mensaje "[IA] GPU: NVIDIA Quadro P1000" la GPU está activa y el sistema '
        'funcionará a máxima velocidad.', 'consejo')

# ── PANTALLA PRINCIPAL ────────────────────────────────────────────────────────
h('3. Pantalla Principal — Vista de Cámaras')
p('Al abrir la aplicación verás la vista de Cámaras. Es la pantalla principal desde donde '
  'puedes ver el video de todas las cámaras configuradas en tiempo real.')
img_from_fig(make_ui_mockup_cameras(), 14,
             'Figura 1. Vista principal de cámaras con grilla 3×2, barra superior y estadísticas')

p('Elementos de la pantalla:', bold=True)
feature_table(
    ['Zona','Descripción'],
    [('Menú lateral izquierdo','Navega entre todas las vistas del sistema (6 opciones).'),
     ('Barra superior','Estado del sistema, selector de modo y gestión de cámaras.'),
     ('Grilla de cámaras','Hasta 6 cámaras en formato 3x2. Las inactivas muestran "SIN VIDEO".'),
     ('Panel inferior (solo aquí)','Estadísticas: frames procesados, FPS, modo activo, cámaras.'),
     ('Boton SYSTEM ONLINE','Indicador de estado del sistema (verde = funcionando).')]
)
callout('El panel de estadísticas ("Frames procesados", "FPS objetivo") solo aparece en '
        'esta vista. Al navegar a otras vistas desaparece para liberar espacio.', 'nota')

# ── MENÚ LATERAL ─────────────────────────────────────────────────────────────
h('4. Menú Lateral — Navegación')
p('Haz clic en cualquier ítem del menú lateral para ir a esa vista. '
  'El ítem seleccionado se resalta en rojo.')
feature_table(
    ['Ítem','Acceso rápido','Para qué sirve'],
    [('Dashboard','Clic en "Dashboard"','Ver cámaras en tiempo real. Igual que "Cámaras".'),
     ('Cámaras','Clic en "Cámaras"','Vista principal con video en vivo.'),
     ('Incidentes','Clic en "Incidentes"','Ver historial de incidentes detectados y exportar a Excel.'),
     ('Destinatarios','Clic en "Destinatarios"','Gestionar quién recibe las alertas automáticas.'),
     ('Notificaciones','Clic en "Notificaciones"','Activar/desactivar Telegram y Gmail.'),
     ('Ajustes','Clic en "Ajustes"','Controlar módulos por cámara y velocidad de procesamiento.')]
)

# ── MODO DE DETECCIÓN ─────────────────────────────────────────────────────────
h('5. Modos de Detección')
p('El selector "Modo" en la barra superior controla qué tipo de análisis se aplica '
  'en tiempo real sobre el video de las cámaras:')
img_from_fig(make_detection_modes(), 14,
             'Figura 2. Los cuatro modos de detección disponibles con sus características')

callout('IMPORTANTE: La detección de incendios está activa SIEMPRE, en TODOS los modos '
        'y en TODAS las cámaras. No necesitas activarla manualmente.', 'advertencia')

feature_table(
    ['Modo','Cuándo usarlo','Consumo de recursos'],
    [('Deteccion General','Vigilancia general sin análisis específico. Para monitor de seguridad básico.','Bajo'),
     ('Deteccion de Accidentes','En cámaras orientadas a intersecciones o vías de tráfico.','Medio'),
     ('Rastreo de Personas','Para identificar y seguir personas específicas por apariencia.','Medio'),
     ('Deteccion de Robos','En tiendas, zonas de acceso controlado o áreas de riesgo. El más intensivo.','Alto')]
)
steps_table([
    (1, 'Ubica el selector "Modo:" en la barra superior.'),
    (2, 'Haz clic en el combo box y selecciona el modo deseado.'),
    (3, 'El cambio es inmediato — no es necesario reiniciar.'),
    (4, 'Si el sistema va lento, selecciona "Detección General" para reducir la carga.'),
])

# ── INCIDENTES ────────────────────────────────────────────────────────────────
h('6. Vista de Incidentes')
p('Aquí puedes ver todos los incidentes que el sistema ha detectado y registrado '
  'automáticamente en la base de datos. Los registros se conservan aunque cierres y vuelvas '
  'a abrir la aplicación.')
img_from_fig(make_ui_mockup_incidents(), 14,
             'Figura 3. Vista de Incidentes con tabla de registros e indicadores de severidad')

feature_table(
    ['Columna','Qué indica'],
    [('Hora','Fecha y hora exacta en que se detectó el incidente.'),
     ('Cámara','Nombre de la cámara donde ocurrió el evento.'),
     ('Tipo','fire = Incendio  |  accident = Choque  |  theft = Robo'),
     ('Severidad','ALTO (confianza > 70%)  |  MEDIO (40-70%)  |  BAJO (< 40%)'),
     ('Confianza','Valor entre 0.0 y 1.0 que indica la certeza del modelo de IA.')]
)

h('6.1. Exportar incidentes a Excel', level=2, color=AZUL2)
steps_table([
    (1, 'Ve a "Incidentes" desde el menú lateral.'),
    (2, 'Haz clic en el botón azul "Exportar a Excel" (esquina superior izquierda).'),
    (3, 'Elige la carpeta de destino y escribe el nombre del archivo.'),
    (4, 'Haz clic en Guardar. El archivo .xlsx se abrirá con Excel o LibreOffice Calc.'),
    (5, 'El archivo incluye todas las columnas: ID, Hora, Cámara, Tipo, Severidad, Confianza y Detalles.'),
])
callout('Los incidentes se guardan en el archivo incidents.db en la carpeta del proyecto. '
        'Se recomienda hacer una copia semanal de este archivo.', 'consejo')

# ── DESTINATARIOS ─────────────────────────────────────────────────────────────
h('7. Vista de Destinatarios')
p('Los destinatarios son las personas que recibirán las alertas automáticas cuando el '
  'sistema detecte un incendio. Puedes tener múltiples destinatarios.')
h('7.1. Agregar un destinatario', level=2, color=AZUL2)
steps_table([
    (1, 'Ve a "Destinatarios" desde el menú lateral.'),
    (2, 'En el campo "Nombre" escribe el nombre completo del destinatario.'),
    (3, 'En "email@example.com" escribe su correo electrónico (opcional).'),
    (4, 'En "+1234567890" escribe su número con código de país (opcional). Ej: +51912345678'),
    (5, 'Haz clic en "Agregar destinatario". Aparecerá en la tabla de abajo.'),
])
h('7.2. Eliminar un destinatario', level=2, color=AZUL2)
steps_table([
    (1, 'En la tabla, haz clic en la fila del destinatario que deseas eliminar.'),
    (2, 'Haz clic en el botón rojo "Eliminar seleccionado".'),
    (3, 'El registro se eliminará permanentemente de la base de datos.'),
])
callout('Los destinatarios se guardan en la base de datos local. Al reiniciar la aplicación '
        'seguirán disponibles sin necesidad de volver a configurarlos.', 'nota')

# ── NOTIFICACIONES ────────────────────────────────────────────────────────────
h('8. Vista de Notificaciones')
p('Desde aquí configuras los canales de comunicación que usará el sistema para enviar '
  'alertas automáticas. Puedes activar o desactivar Telegram y Gmail de forma independiente.')
img_from_fig(make_ui_mockup_notif(), 14,
             'Figura 4. Vista de Notificaciones con configuración de Telegram y Gmail')

h('8.1. Configurar Telegram', level=2, color=AZUL2)
steps_table([
    (1, 'En la sección "Telegram", marca la casilla "Habilitado".'),
    (2, 'Obtén el Token del bot: busca @BotFather en Telegram → /newbot → copia el token.'),
    (3, 'Pega el token en el campo "Token del bot".'),
    (4, 'Obtén tu Chat ID: escribe /start a tu bot, luego visita: https://api.telegram.org/botTU_TOKEN/getUpdates'),
    (5, 'Copia el número "id" y pégalo en "Chat ID(s)". Para múltiples chats separa con coma.'),
    (6, 'Haz clic en "Guardar configuracion de notificaciones".'),
    (7, 'Para desactivar: desmarca "Habilitado" y guarda.'),
])
callout('Las alertas de Telegram incluyen texto con el tipo de incidente, cámara y hora. '
        'Si el sistema tiene acceso a internet, también adjunta una imagen del incidente.', 'nota')

h('8.2. Configurar Gmail', level=2, color=AZUL2)
steps_table([
    (1, 'En la sección "Gmail", marca la casilla "Habilitado".'),
    (2, 'Escribe tu correo de Gmail en "Correo remitente".'),
    (3, 'Genera una Contraseña de Aplicación: myaccount.google.com → Seguridad → Contraseñas de app.'),
    (4, 'Copia los 16 caracteres y pégalos en "Contrasena de aplicacion".'),
    (5, 'Escribe los correos destinatarios separados por coma en "Destinatarios".'),
    (6, 'Haz clic en "Guardar configuracion de notificaciones".'),
])
callout('ADVERTENCIA: Usa una Contraseña de Aplicación (16 caracteres generada por Google), '
        'no tu contraseña normal de Gmail. Requiere tener activada la verificación en 2 pasos.', 'advertencia')

# ── AJUSTES ────────────────────────────────────────────────────────────────────
h('9. Vista de Ajustes — Control de Rendimiento')
p('Esta vista te permite controlar exactamente qué módulos están activos en cada cámara, '
  'lo que te permite gestionar el consumo de CPU y GPU del sistema según tus necesidades.')
img_from_fig(make_ui_mockup_ajustes(), 14,
             'Figura 5. Vista de Ajustes con módulos por cámara y control de FPS')

h('9.1. Habilitar/Deshabilitar módulos por cámara', level=2, color=AZUL2)
p('Para cada cámara configurada verás un panel con tres casillas:')
feature_table(
    ['Módulo','Desactívalo si...','Consumo de GPU'],
    [('Incendios','La cámara está en una zona sin riesgo de fuego (ej: sala de reuniones).','Bajo'),
     ('Robos','La cámara está en un área pública sin objetos de valor o solo con acceso controlado.','Alto'),
     ('Choques','La cámara no apunta a vías de tráfico vehicular.','Medio')]
)
steps_table([
    (1, 'Ve a "Ajustes" desde el menú lateral.'),
    (2, 'Localiza el panel de la cámara que deseas configurar (identificada por su nombre).'),
    (3, 'Marca o desmarca las casillas de los módulos según necesites.'),
    (4, 'Haz clic en "Guardar ajustes". Los cambios se aplican INMEDIATAMENTE sin reiniciar.'),
])
callout('Recomendación: Si el sistema va lento (FPS < 5), desactiva el módulo "Robos" en '
        'las cámaras donde no sea necesario. Es el módulo que más recursos consume.', 'consejo')

h('9.2. Ajustar el FPS objetivo', level=2, color=AZUL2)
p('El campo "FPS objetivo" controla la velocidad de procesamiento del sistema:')
feature_table(
    ['FPS configurado','Uso de GPU/CPU','Cuándo usarlo'],
    [('5 - 10 FPS','Muy bajo (< 30%)','Cámaras de edificios, archivos, zonas de poco movimiento.'),
     ('15 - 20 FPS','Medio (40-60%)','Uso diario normal. Buen equilibrio entre calidad y recursos.'),
     ('25 FPS (por defecto)','Alto (60-80%)','Zonas de alta actividad o cuando la GPU es potente.'),
     ('30 FPS','Máximo (GPU requerida)','Solo con GPU dedicada de 4+ GB VRAM.')]
)
callout('El valor de FPS configura el tiempo entre cada análisis de video. '
        'El video siempre se captura de forma continua; solo varía la frecuencia del análisis.', 'nota')

# ── GESTIÓN DE CÁMARAS ────────────────────────────────────────────────────────
h('10. Gestión de Cámaras RTSP')
p('Para agregar, editar o eliminar cámaras haz clic en "Gestionar Camaras" en la barra superior.')
img_from_fig(make_camera_dialog(), 12,
             'Figura 6. Dialogo de gestión de cámaras con formulario de alta y lista de configuradas')

h('10.1. Agregar una cámara', level=2, color=AZUL2)
steps_table([
    (1, 'Haz clic en el boton rojo "Gestionar Camaras" en la barra superior.'),
    (2, 'En el campo "Nombre de la camara" escribe un nombre descriptivo (ej: "Entrada Principal").'),
    (3, 'En el campo de URL escribe la dirección RTSP de la cámara.'),
    (4, 'Formato: rtsp://USUARIO:CONTRASEÑA@IP:PUERTO/stream   (ej: rtsp://admin:1234@192.168.1.100:554/stream)'),
    (5, 'Haz clic en "+ Agregar". La cámara aparecerá en la lista.'),
    (6, 'Cierra el diálogo. La cámara aparecerá en la grilla principal.'),
])
callout('Puedes usar "Probar conexion" para verificar que la URL RTSP sea correcta '
        'antes de guardar. Si no se puede conectar, revisa la IP y credenciales.', 'consejo')

h('10.2. Formatos de URL RTSP comunes', level=2, color=AZUL2)
feature_table(
    ['Marca','Formato URL típico'],
    [('Hikvision','rtsp://admin:contraseña@192.168.X.X:554/Streaming/Channels/101'),
     ('Dahua','rtsp://admin:contraseña@192.168.X.X:554/cam/realmonitor?channel=1&subtype=0'),
     ('TP-Link Tapo','rtsp://usuario:contraseña@192.168.X.X:554/stream1'),
     ('Genérica','rtsp://admin:contraseña@192.168.X.X:554/'),
     ('Archivo de video','C:\\Videos\\grabacion.mp4  (ruta completa al archivo)')]
)

# ── MODELO DE DETECCIÓN ───────────────────────────────────────────────────────
h('11. Cómo Funciona la Detección (Información Técnica)')
p('Los tres módulos de detección usan modelos de inteligencia artificial '
  'entrenados específicamente para cada tipo de incidente:')

img_file(os.path.join(ROOT,'modulos','incendio','results.png'), 13,
         'Figura 7. Curvas de entrenamiento del modelo YOLOv8 para detección de incendios/humo.\n'
         'Arriba: pérdidas de entrenamiento. Abajo: métricas de validación (precisión, recall, mAP).')

feature_table(
    ['Módulo','Modelo','Dataset de entrenamiento','Precisión estimada'],
    [('Incendios','incendio/best.pt (YOLOv8)','Fire-8: 1200+ imágenes de fuego/humo','87% en condiciones normales'),
     ('Choques','choques/best.pt (YOLOv8)','Dataset de accidentes de tráfico','78% en escenas viales'),
     ('Robos','yolov8n.pt + lógica propia','COCO (personas y objetos) + reglas','73% en detección de transferencia')]
)

callout('Los modelos se ejecutan completamente en tu equipo local. Los videos de las cámaras '
        'NUNCA se envían a servidores externos. Solo las alertas de texto viajan a Telegram/Gmail.', 'consejo')

# ── FAQ ───────────────────────────────────────────────────────────────────────
h('12. Solución de Problemas Frecuentes')

problemas=[
    ('La cámara muestra "SIN VIDEO"',
     'Verifica: (1) La URL RTSP en "Gestionar Cámaras" es correcta. '
     '(2) La cámara está encendida y conectada a la misma red. '
     '(3) El usuario/contraseña en la URL son correctos. '
     '(4) No hay firewall bloqueando el puerto 554.'),
    ('La aplicación tarda mucho en abrir (> 30 segundos)',
     'Esto es normal la primera vez. El sistema debe compilar los modelos de IA '
     'para la GPU. En los siguientes arranques será más rápido (10-15 segundos).'),
    ('No recibo alertas de Telegram',
     '(1) Ve a Notificaciones y verifica que "Habilitado" esté marcado. '
     '(2) Revisa que el Token del bot sea correcto. '
     '(3) Asegúrate de haber escrito /start a tu bot en Telegram. '
     '(4) Verifica que el Chat ID sea el correcto (no confundir con el ID del bot).'),
    ('No recibo alertas de Gmail',
     '(1) Usa una Contraseña de Aplicación, no tu contraseña de Gmail. '
     '(2) Activa la verificación en 2 pasos en tu cuenta Google. '
     '(3) Genera la contraseña en: myaccount.google.com → Seguridad → Contraseñas de app. '
     '(4) El remitente y el correo autenticado deben ser la misma cuenta.'),
    ('El sistema va muy lento (menos de 5 FPS)',
     '(1) Ve a Ajustes y baja el FPS objetivo a 10. '
     '(2) Desactiva el módulo "Robos" en las cámaras que no lo necesiten. '
     '(3) Cierra otras aplicaciones que usen la GPU. '
     '(4) Usa el modo "Detección General" que consume menos recursos.'),
    ('El sistema detecta demasiados falsos positivos de incendio',
     'Puede ocurrir con luces de neón, reflejos solares o iluminación roja. '
     'Solución: Ve al archivo fire_config.json y aumenta el valor de "threshold" de 0.12 a 0.25. '
     'Esto hace el detector más estricto.'),
    ('¿Cómo sé si una detección es confiable?',
     'Observa la columna "Confianza" en la vista de Incidentes: '
     'Confianza > 0.7 = Detección muy confiable. '
     'Confianza 0.4-0.7 = Posible detección, verificar manualmente. '
     'Confianza < 0.4 = Puede ser un falso positivo.'),
]
for pregunta, respuesta in problemas:
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'
    cell=t.rows[0].cells[0]; shd(cell,'EBF3FB')
    pq=cell.paragraphs[0]
    r1=pq.add_run('P: '+pregunta+'\n')
    r1.bold=True; r1.font.size=Pt(10.5)
    r1.font.color.rgb=AZUL2
    r2=pq.add_run('R: '+respuesta)
    r2.font.size=Pt(10)
    doc.add_paragraph()

# ── REFERENCIA RÁPIDA ─────────────────────────────────────────────────────────
h('13. Guía de Referencia Rápida')
p('Tarjeta de referencia para el personal operativo del serenazgo:', bold=True)
t=doc.add_table(rows=9,cols=2); t.style='Table Grid'
for i,(accion,donde) in enumerate([
    ('ACCIÓN','DÓNDE / CÓMO'),
    ('Ver cámaras en vivo','Menú → Dashboard o Cámaras'),
    ('Ver historial de incidentes','Menú → Incidentes'),
    ('Exportar incidentes a Excel','Menú → Incidentes → botón "Exportar a Excel"'),
    ('Agregar destinatario de alertas','Menú → Destinatarios → rellenar formulario'),
    ('Activar/desactivar Telegram','Menú → Notificaciones → casilla "Habilitado"'),
    ('Desactivar módulo por cámara','Menú → Ajustes → desmarcar casilla de la cámara'),
    ('Agregar nueva cámara','Barra superior → botón "Gestionar Cámaras"'),
    ('Cambiar modo de detección','Barra superior → selector "Modo"'),
]):
    t.rows[i].cells[0].text=accion; t.rows[i].cells[1].text=donde
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size=Pt(10)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size=Pt(10)
    if i==0:
        for j in range(2):
            t.rows[i].cells[j].paragraphs[0].runs[0].bold=True
            t.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb=BLANCO
            shd(t.rows[i].cells[j],'1F2A5E')
    elif i%2==0: shd(t.rows[i].cells[0],'F0F5FB'); shd(t.rows[i].cells[1],'F0F5FB')

doc.add_paragraph()
callout('Para soporte técnico contactar a: Yords Williams Ccalla Mamani — '
        'Tel: +51 930 240 476 — Practicante de Ingeniería de Software y Sistemas, UNAJ.', 'nota')

# ── FIRMA ───────────────────────────────────────────────────────────────────
doc.add_page_break()
c('Juliaca, julio del 2026', size=11)
doc.add_paragraph()
c('_________________________________')
c('Yords Williams Ccalla Mamani', bold=True, size=12)
c('Practicante — Municipalidad Distrital de Caracoto', size=10)
c('E.P. Ingeniería de Software y Sistemas — UNAJ', size=10)

out=os.path.join(ROOT,'MANUAL_USUARIO_v2.docx')
doc.save(out)
print(f'✓ Guardado: {out}')
