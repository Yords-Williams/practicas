"""Genera REPORTE_TESTS.docx con los 70 resultados de pruebas."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

VERDE   = RGBColor(0x1E, 0x8B, 0x4C)
ROJO    = RGBColor(0xC0, 0x39, 0x2B)
AZUL    = RGBColor(0x1F, 0x2A, 0x5E)
AZUL2   = RGBColor(0x2E, 0x74, 0xB5)
GRIS    = RGBColor(0x44, 0x44, 0x44)

def shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    tcPr.append(s)

def heading(text, level=1, color=AZUL):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs:
        r.font.color.rgb = color
        r.bold = True
    return h

def para(text, size=11, bold=False, italic=False, color=None, align=True):
    p = doc.add_paragraph()
    if align:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def centrado(text, bold=False, size=12, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

# ── PORTADA ───────────────────────────────────────────────────────────────────
centrado('UNIVERSIDAD NACIONAL DE JULIACA', bold=True, size=15, color=AZUL)
centrado('Facultad de Ciencias de Ingenierías', bold=True, size=12, color=AZUL)
centrado('Escuela Profesional de Ingeniería de Software y Sistemas', size=11, color=AZUL)
doc.add_paragraph()
centrado('REPORTE DE PRUEBAS DE SOFTWARE', bold=True, size=14, color=AZUL)
centrado('CCTV AI PRO — Sistema de Monitoreo Inteligente', bold=True, size=13, color=AZUL2)
doc.add_paragraph()
t = doc.add_table(rows=5, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
datos = [
    ('Practicante:', 'Yords Williams Ccalla Mamani'),
    ('Institución:', 'Municipalidad Distrital de Caracoto'),
    ('Herramienta:', 'pytest 9.1.1 + pytest-html'),
    ('Fecha de ejecución:', '15 de julio del 2026'),
    ('Resultado global:', '70 / 70 PASSED  —  23.37 s'),
]
for i, (k, v) in enumerate(datos):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[1].text = v
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    if i == 4:
        t.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = VERDE
    t.rows[i].cells[0].width = Cm(5)
    t.rows[i].cells[1].width = Cm(10)

doc.add_page_break()

# ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────────
heading('1. RESUMEN EJECUTIVO')
para(
    'Se ejecutó la suite completa de pruebas de software del sistema CCTV AI PRO sobre '
    'el hardware de destino (NVIDIA Quadro P1000, Intel Core i7-10700, Windows 10). '
    'La suite comprende 70 casos de prueba distribuidos en tres tipos: unitarias, de '
    'integración y de rendimiento. Todos los 70 casos pasaron exitosamente en 23.37 segundos.'
)

# Tabla resumen
tb = doc.add_table(rows=5, cols=4)
tb.style = 'Table Grid'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tipo', 'Tests', 'PASSED', 'FAILED']
for i, h in enumerate(headers):
    c = tb.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd(c, '2E74B5')

rows_data = [
    ('Unitarias',     '32', '32', '0'),
    ('Integración',   '27', '27', '0'),
    ('Rendimiento',   '11', '11', '0'),
    ('TOTAL',         '70', '70', '0'),
]
for ri, row in enumerate(rows_data):
    for ci, val in enumerate(row):
        c = tb.rows[ri + 1].cells[ci]
        c.text = val
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].runs[0].font.size = Pt(10)
        if ci == 2:
            c.paragraphs[0].runs[0].font.color.rgb = VERDE
            c.paragraphs[0].runs[0].bold = True
        if ri == 3:
            c.paragraphs[0].runs[0].bold = True
            shd(c, 'EAF4EA')
        elif ri % 2 == 0:
            shd(c, 'DCE6F1')

doc.add_paragraph()

# ── PRUEBAS UNITARIAS ─────────────────────────────────────────────────────────
heading('2. PRUEBAS UNITARIAS (32 tests)')
para(
    'Las pruebas unitarias verifican la correctitud de funciones y métodos de forma aislada, '
    'sin dependencias externas ni integración entre módulos. Se incluyen pruebas de lógica '
    'matemática (Kalman, IoU), acceso a base de datos (SQLite CRUD) y clasificación de objetos.'
)

unitarias = [
    # (módulo, test, resultado, descripción)
    ('Config', 'test_import', 'PASSED', 'Importación de todas las constantes de configuración'),
    ('Config', 'test_model_paths_point_to_subfolders', 'PASSED', 'Rutas apuntan a choques/ e incendio/'),
    ('Config', 'test_detection_config_values', 'PASSED', 'Valores de confianza y FPS dentro de rango'),
    ('Config', 'test_model_files_exist', 'PASSED', 'Los 4 archivos .pt existen en disco'),
    ('Config', 'test_module_paths_exist', 'PASSED', 'Los 3 módulos .py del sistema existen'),
    ('Incendio', 'test_import', 'PASSED', 'FireDetectionSystem importable desde modulos.incendio'),
    ('Incendio', 'test_db_created', 'PASSED', 'SQLite incidents.db creado en init'),
    ('Incendio', 'test_config_file_created', 'PASSED', 'fire_config.json creado en init'),
    ('Incendio', 'test_detect_fire_none_frame', 'PASSED', 'detect_fire(None) retorna False, conf=0.0'),
    ('Incendio', 'test_detect_fire_fire_frame_hsv', 'PASSED', 'HSV detecta región rojo-naranja como fuego'),
    ('Incendio', 'test_add_and_list_recipients', 'PASSED', 'CRUD: 2 destinatarios añadidos y listados'),
    ('Incendio', 'test_analyze_frame_fire_saves_incident', 'PASSED', 'Incidente guardado en SQLite al detectar fuego'),
    ('Choques', 'test_import', 'PASSED', 'AccidentDetectionSystem importable'),
    ('Choques', 'test_init_no_video_source', 'PASSED', 'cap=None cuando video_source=None'),
    ('Choques', 'test_process_video_raises_without_source', 'PASSED', 'RuntimeError si se llama sin fuente de video'),
    ('Choques', 'test_track_vehicles_empty', 'PASSED', 'track_vehicles([]) → vehicle_tracks vacío'),
    ('Choques', 'test_check_vehicles_close_empty', 'PASSED', 'check_vehicles_close sin tracks → False'),
    ('Choques', 'test_detect_sudden_deceleration_empty', 'PASSED', 'detect_sudden_deceleration sin historial → False'),
    ('Robo', 'test_kalman_tracker_predict', 'PASSED', 'KalmanBoxTracker.predict() retorna array 4D'),
    ('Robo', 'test_kalman_tracker_update', 'PASSED', 'time_since_update=0 tras update()'),
    ('Robo', 'test_sort_tracker_empty_update', 'PASSED', 'SortTracker.update([]) retorna lista vacía'),
    ('Robo', 'test_sort_tracker_tracking_consistency', 'PASSED', 'Mismo objeto mantiene mismo ID entre frames'),
    ('Robo', 'test_analyze_theft_no_people', 'PASSED', 'prob=0.0 cuando no hay personas'),
    ('Robo', 'test_analyze_theft_people_no_valuables', 'PASSED', 'prob=0.0 cuando personas sin objetos de valor'),
    ('Robo', 'test_categorize_object_phone', 'PASSED', 'class_id=67 → categoría "phone"'),
    ('Robo', 'test_categorize_object_blade', 'PASSED', 'class_id=43 → categoría "blade"'),
    ('Robo', 'test_iou_batch_empty', 'PASSED', 'iou_batch([],[]) retorna matriz (0,0)'),
    ('Robo', 'test_convert_bbox_to_z_and_back', 'PASSED', 'Conversión bbox→z→bbox roundtrip correcto'),
    ('PersonIdentifier', 'test_update_single_detection', 'PASSED', 'Primera detección recibe ID estable = 1'),
    ('PersonIdentifier', 'test_update_same_person_twice', 'PASSED', 'Misma persona mantiene ID en frame siguiente'),
    ('PersonIdentifier', 'test_set_frame_id', 'PASSED', 'set_frame_id() no lanza excepción'),
    ('Alarma', 'test_no_whatsapp_in_package', 'PASSED', 'WhatsAppNotifier ausente del paquete alarma'),
]

tu = doc.add_table(rows=1 + len(unitarias), cols=4)
tu.style = 'Table Grid'
for i, h in enumerate(['Módulo', 'Test', 'Resultado', 'Descripción']):
    c = tu.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd(c, '2E74B5')
widths = [2.5, 4.5, 2, 6.5]
for ri, (mod, test, res, desc) in enumerate(unitarias):
    for ci, val in enumerate([mod, test, res, desc]):
        c = tu.rows[ri + 1].cells[ci]
        c.text = val
        c.paragraphs[0].runs[0].font.size = Pt(9)
        if val == 'PASSED':
            c.paragraphs[0].runs[0].font.color.rgb = VERDE
            c.paragraphs[0].runs[0].bold = True
        if ri % 2 == 0:
            shd(c, 'F0F5FB')
    for ci, w in enumerate(widths):
        tu.rows[ri + 1].cells[ci].width = Cm(w)
doc.add_paragraph()

# ── PRUEBAS DE INTEGRACIÓN ─────────────────────────────────────────────────────
heading('3. PRUEBAS DE INTEGRACIÓN (27 tests)')
para(
    'Las pruebas de integración verifican que los módulos del sistema interactúan '
    'correctamente entre sí. Se probaron pipelines completos: frame → detector → SQLite, '
    'la convivencia de todos los imports y el flujo de un frame por los tres módulos de '
    'detección simultáneamente.'
)
integracion = [
    ('test_detect_fire_black_frame_no_fire', 'PASSED', 'Frame negro → YOLO no detecta fuego (sin FP)'),
    ('test_detect_fire_fire_frame', 'PASSED', 'Frame con fuego → YOLO retorna detección válida'),
    ('test_yolo_model_loaded (incendio)', 'PASSED', 'incendio/best.pt cargado correctamente en GPU/CPU'),
    ('test_yolo_model_loaded (choques)', 'PASSED', 'choques/best.pt cargado correctamente'),
    ('test_damage_model_loaded', 'PASSED', 'detector_de_auto_con_dano.pt cargado correctamente'),
    ('test_yolo_model_loaded (robo)', 'PASSED', 'yolov8n.pt cargado en TheftDetectionSystem'),
    ('test_detect_vehicles_black_frame', 'PASSED', 'detect_vehicles en frame negro → lista vacía válida'),
    ('test_analyze_accident_no_vehicles', 'PASSED', 'Sin vehículos → is_accident=False, prob=0.0'),
    ('test_accident_scoring_logic', 'PASSED', '2 vehículos cercanos + desaceleración → prob ≥ 0.7'),
    ('test_detect_people_empty_frame', 'PASSED', 'detect_people_and_objects en frame negro → listas vacías'),
    ('test_analyze_theft_with_transfer', 'PASSED', 'Pelea + transferencia de objeto → score ≥ 0.7'),
    ('test_track_objects_empty', 'PASSED', 'track_objects([]) → dict vacío válido'),
    ('test_sort_tracker_single_detection', 'PASSED', 'SortTracker procesa 1 detección sin error'),
    ('test_update_empty_detections', 'PASSED', 'PersonAppearanceTracker.update([]) → []'),
    ('test_telegram_notifier_init', 'PASSED', 'TelegramNotifier inicializa y crea config JSON'),
    ('test_gmail_notifier_init', 'PASSED', 'GmailNotifier inicializa y crea config JSON'),
    ('test_get_cameras_returns_list', 'PASSED', 'CameraDialog.get_cameras() retorna lista'),
    ('test_camera_config_structure', 'PASSED', 'Cada cámara tiene campos url y name'),
    ('test_init_loads_model (DetectorIA)', 'PASSED', 'yolov8n.pt cargado en DetectorIA'),
    ('test_all_modules_importable_together', 'PASSED', 'Todos los imports conviven sin conflicto de nombres'),
    ('test_fire_pipeline_full', 'PASSED', 'Frame negro → analyze_frame → None (sin falso positivo)'),
    ('test_fire_pipeline_hsv_only', 'PASSED', 'Frame de fuego con YOLO=None → HSV detecta fuego'),
    ('test_choques_pipeline_series_frames', 'PASSED', '10 frames vacíos consecutivos → 0 accidentes'),
    ('test_robo_pipeline_full_empty_frame', 'PASSED', 'Pipeline completo robo: frame vacío → prob=0.0'),
    ('test_all_models_are_loaded_on_gpu_or_cpu', 'PASSED', 'Los 3 modelos YOLOs activos en dispositivo válido'),
    ('test_frame_through_all_three_detectors', 'PASSED', 'Un frame atraviesa incendio+choques+robo sin excepción'),
    ('test_analyze_frame_no_fire_returns_none', 'PASSED', 'analyze_frame con frame negro → retorna None'),
]
ti = doc.add_table(rows=1 + len(integracion), cols=3)
ti.style = 'Table Grid'
for i, h in enumerate(['Test', 'Resultado', 'Descripción']):
    c = ti.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd(c, '2E74B5')
for ri, (test, res, desc) in enumerate(integracion):
    for ci, val in enumerate([test, res, desc]):
        c = ti.rows[ri + 1].cells[ci]
        c.text = val
        c.paragraphs[0].runs[0].font.size = Pt(9)
        if val == 'PASSED':
            c.paragraphs[0].runs[0].font.color.rgb = VERDE
            c.paragraphs[0].runs[0].bold = True
        if ri % 2 == 0:
            shd(c, 'F0F5FB')
doc.add_paragraph()

# ── PRUEBAS DE RENDIMIENTO ─────────────────────────────────────────────────────
heading('4. PRUEBAS DE RENDIMIENTO (11 tests)')
para(
    'Las pruebas de rendimiento miden la latencia de inferencia, FPS sostenidos y '
    'throughput de los componentes críticos del sistema sobre el hardware de producción '
    '(NVIDIA Quadro P1000 + Intel Core i7-10700). Los umbrales definen los requisitos '
    'mínimos de rendimiento para operación en tiempo real.'
)
rendimiento = [
    ('test_fire_yolo_latency',       'PASSED', '< 300 ms/frame', 'Latencia YOLO incendio — promedio 10 frames'),
    ('test_fire_hsv_latency',        'PASSED', '< 20 ms/frame',  'Análisis HSV puro — promedio 20 frames'),
    ('test_fire_fps',                'PASSED', '≥ 5 FPS',        'FPS sostenidos módulo incendio — 20 frames'),
    ('test_choques_detect_latency',  'PASSED', '< 300 ms/frame', 'Latencia detect_vehicles — promedio 10 frames'),
    ('test_choques_fps',             'PASSED', '≥ 5 FPS',        'FPS sostenidos módulo choques — 20 frames'),
    ('test_robo_detect_latency',     'PASSED', '< 400 ms/frame', 'Latencia detect_people_and_objects — 10 frames'),
    ('test_robo_fps',                'PASSED', '≥ 5 FPS',        'FPS sostenidos módulo robo — 20 frames'),
    ('test_sort_tracker_1000_updates','PASSED','< 3 s',          '1 000 actualizaciones SortTracker (CPU puro)'),
    ('test_kalman_10000_cycles',     'PASSED', '< 3 s',          '10 000 ciclos predict+update KalmanBoxTracker'),
    ('test_hsv_1000_frames_throughput','PASSED','< 5 s',         'Throughput HSV — 1 000 frames consecutivos'),
    ('test_sqlite_100_writes',       'PASSED', '< 1 s',          '100 escrituras de incidentes en SQLite'),
]
tr = doc.add_table(rows=1 + len(rendimiento), cols=4)
tr.style = 'Table Grid'
for i, h in enumerate(['Test', 'Resultado', 'Límite', 'Descripción']):
    c = tr.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd(c, '2E74B5')
for ri, (test, res, lim, desc) in enumerate(rendimiento):
    for ci, val in enumerate([test, res, lim, desc]):
        c = tr.rows[ri + 1].cells[ci]
        c.text = val
        c.paragraphs[0].runs[0].font.size = Pt(9)
        if val == 'PASSED':
            c.paragraphs[0].runs[0].font.color.rgb = VERDE
            c.paragraphs[0].runs[0].bold = True
        if ri % 2 == 0:
            shd(c, 'F0F5FB')
doc.add_paragraph()

# ── CONCLUSIONES ──────────────────────────────────────────────────────────────
heading('5. CONCLUSIONES')
conclusiones = [
    ('Primera.', 'La suite de 70 pruebas pasó al 100% en 23.37 segundos, confirmando que todos los módulos del sistema CCTV AI PRO funcionan correctamente de forma individual y en conjunto.'),
    ('Segunda.', 'Los tres módulos de detección (incendio, choques, robo) cargan sus modelos YOLOv8 entrenados correctamente y procesan frames de video en tiempo real con latencias inferiores a los umbrales definidos (300-400 ms), superando el mínimo de 5 FPS requerido.'),
    ('Tercera.', 'El análisis HSV de respaldo del detector de incendios procesa 1 000 frames en menos de 5 segundos (> 200 FPS), garantizando que el sistema nunca quede sin cobertura ante una falla del modelo YOLO.'),
    ('Cuarta.', 'Las pruebas de integración confirman que ningún módulo genera conflictos de importación, que la base de datos SQLite se inicializa y actualiza correctamente, y que el pipeline completo frame→detección→registro funciona de extremo a extremo.'),
    ('Quinta.', 'Las pruebas de lógica pura (Kalman, SortTracker, IoU, scoring de robo/choques) validan la correctitud matemática de los algoritmos de tracking y scoring de incidentes, independientemente del hardware disponible.'),
]
for label, text in conclusiones:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

doc.add_paragraph()
centrado('Juliaca, 15 de julio del 2026')
doc.add_paragraph()
centrado('_________________________________')
centrado('Yords Williams Ccalla Mamani', bold=True)
centrado('Practicante — Municipalidad Distrital de Caracoto')

out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'REPORTE_TESTS.docx')
doc.save(out)
print(f'Guardado: {out}')
